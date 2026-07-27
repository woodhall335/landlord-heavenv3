"""Build branded tenancy HBS templates from the pinned official model DOCX files.

The government documents are legal-reference inputs only. Customer PDFs continue
to use Landlord Heaven's HTML/PDF renderer and visual language.
"""

from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCES = ROOT / "config/mqs/tenancy_agreement/official_model_sources"
WALES_TEMPLATES = ROOT / "config/jurisdictions/uk/wales/templates"
SCOTLAND_TEMPLATES = ROOT / "config/jurisdictions/uk/scotland/templates"


CSS = """
@page { size: A4; margin: 22mm 18mm 22mm; }
* { box-sizing: border-box; }
body { margin: 0; color: #172033; font: 10.2pt/1.46 Georgia, "Times New Roman", serif; }
h1, h2, h3 { font-family: Arial, Helvetica, sans-serif; break-after: avoid; }
h1 { font-size: 21pt; letter-spacing: .35pt; text-align: center; text-transform: uppercase; }
h2 { margin: 20pt 0 9pt; padding-bottom: 4pt; border-bottom: 1.2pt solid #243b53; font-size: 13pt; }
h3 { margin: 13pt 0 6pt; font-size: 10.8pt; }
p { margin: 0 0 7pt; orphans: 3; widows: 3; text-align: justify; }
.cover { min-height: 205mm; display: flex; flex-direction: column; justify-content: center; text-align: center; }
.brand { color: #176b57; font: 700 11pt Arial, sans-serif; letter-spacing: 1.5pt; text-transform: uppercase; }
.subtitle { color: #52616b; font: 11pt Arial, sans-serif; text-align: center; }
.legal-source { margin: 18pt auto; max-width: 135mm; padding: 10pt; border: 1pt solid #9fb3c8; background: #f4f8fb; font-size: 9pt; text-align: left; }
.page-break { break-before: page; page-break-before: always; }
.section { margin: 12pt 0; }
.data-table, .signature-table { width: 100%; margin: 9pt 0 14pt; border-collapse: collapse; break-inside: avoid; }
.data-table th, .data-table td, .signature-table th, .signature-table td { border: .65pt solid #829ab1; padding: 6pt 7pt; vertical-align: top; }
.data-table th { width: 34%; background: #edf3f7; font: 700 8.6pt Arial, sans-serif; text-align: left; text-transform: uppercase; }
.tenant-card { margin: 8pt 0; padding: 8pt; border: .65pt solid #bcccdc; break-inside: avoid; }
.term-heading { margin-top: 14pt; color: #102a43; }
.term-number { display: inline-block; min-width: 24pt; }
.legal-paragraph { margin-left: 18pt; }
.legal-list { margin: 2pt 0 8pt 34pt; }
.legal-list-item { margin: 2pt 0; text-align: justify; }
.explanatory { padding: 8pt 10pt; border-left: 3pt solid #176b57; background: #f2f8f6; }
.omitted { padding: 7pt 9pt; border: 1pt solid #c6d4df; background: #f8fafc; font-style: italic; }
.signature-table td { height: 38pt; }
.signature-table .line { height: 54pt; }
.footer-note { margin-top: 24pt; padding-top: 8pt; border-top: .5pt solid #9fb3c8; color: #52616b; font-size: 8.5pt; }
"""


def run_html(paragraph) -> str:
    chunks: list[str] = []
    for run in paragraph.runs:
        text = escape(run.text).replace("\n", "<br>")
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        chunks.append(text)
    return "".join(chunks) or escape(paragraph.text)


def render_range(
    document: Document,
    start: int,
    end: int,
    *,
    heading_counter: int | None = None,
    fixed_wales_terms: bool = False,
    replacements: dict[int, str] | None = None,
    skip: set[int] | None = None,
) -> tuple[str, int | None]:
    output: list[str] = []
    list_open = False
    replacements = replacements or {}
    skip = skip or set()

    def close_list() -> None:
        nonlocal list_open
        if list_open:
            output.append("</ul>")
            list_open = False

    for index in range(start, min(end, len(document.paragraphs))):
        if index in skip:
            continue
        paragraph = document.paragraphs[index]
        text = " ".join(paragraph.text.split())
        if not text:
            close_list()
            continue
        if index in replacements:
            close_list()
            output.append(replacements[index])
            continue

        style = paragraph.style.name
        html = run_html(paragraph)
        is_list = style in {"List Paragraph", "Heading 3", "N1", "N2", "N3", "N5"}

        if fixed_wales_terms and style == "H1":
            close_list()
            assert heading_counter is not None
            output.append(
                f'<h3 class="term-heading"><span class="term-number">{heading_counter}.</span>{html}</h3>'
            )
            heading_counter += 1
        elif heading_counter is not None and not fixed_wales_terms and style in {"Heading 1", "Heading 2"}:
            close_list()
            output.append(
                f'<h3 class="term-heading"><span class="term-number">{heading_counter}.</span>{html}</h3>'
            )
            heading_counter += 1
        elif style in {"Part", "PartHead", "ScheduleHead", "Heading 1", "Heading 2"}:
            close_list()
            output.append(f"<h2>{html}</h2>")
        elif style == "H1":
            close_list()
            output.append(f'<h3 class="term-heading">{html}</h3>')
        elif is_list:
            if not list_open:
                output.append('<ul class="legal-list">')
                list_open = True
            output.append(f'<li class="legal-list-item">{html}</li>')
        else:
            close_list()
            output.append(f'<p class="legal-paragraph">{html}</p>')

    close_list()
    return "\n".join(output), heading_counter


def shell(title: str, subtitle: str, body: str, source_note: str) -> str:
    return f"""<!doctype html>
<html lang="en-GB"><head><meta charset="utf-8"><title>{escape(title)}</title>
<style>{CSS}</style></head><body>
<section class="cover">
  <div class="brand">Landlord Heaven</div>
  <h1>{escape(title)}</h1>
  <p class="subtitle">{escape(subtitle)}</p>
  <div class="legal-source">{source_note}</div>
  <p class="subtitle">Document ID: {{{{document_id}}}}<br>Generated: {{{{current_date}}}}</p>
</section>
{body}
<p class="footer-note">This document should be reviewed in full before signature. Keep the signed agreement, inventory and all prescribed information together.</p>
</body></html>"""


def write_template(path: Path, content: str) -> None:
    normalized = "\n".join(line.rstrip() for line in content.splitlines()) + "\n"
    path.write_text(normalized, encoding="utf-8", newline="\n")


def tenant_rows(label: str) -> str:
    return f"""
<h2>{label}</h2>
{{{{#each tenants}}}}
<div class="tenant-card">
  <strong>{{{{full_name}}}}</strong><br>
  {{{{#if address}}}}{{{{address}}}}<br>{{{{/if}}}}
  Email: {{{{email}}}} &nbsp; Telephone: {{{{phone}}}}
</div>
{{{{/each}}}}"""


def signature_block(occupier_label: str) -> str:
    return f"""
<section class="page-break">
  <h2>Execution</h2>
  <p>By signing, each party confirms that they have read and agree to this written agreement.</p>
  {{{{#each tenants}}}}
  <table class="signature-table">
    <tr><th colspan="2">{occupier_label}: {{{{full_name}}}}</th></tr>
    <tr><td class="line">Signature</td><td>Date</td></tr>
  </table>
  {{{{/each}}}}
  <table class="signature-table">
    <tr><th colspan="2">Landlord: {{{{landlord_full_name}}}}</th></tr>
    <tr><td class="line">Signature</td><td>Date</td></tr>
  </table>
  {{{{#if landlord_2_full_name}}}}
  <table class="signature-table">
    <tr><th colspan="2">Joint landlord: {{{{landlord_2_full_name}}}}</th></tr>
    <tr><td class="line">Signature</td><td>Date</td></tr>
  </table>
  {{{{/if}}}}
</section>"""


def wales_key_matters(fixed: bool) -> str:
    fixed_rows = """
  <tr><th>Fixed term</th><td>{{term_length}}</td></tr>
  <tr><th>Contract ends</th><td>{{format_date tenancy_end_date}}</td></tr>""" if fixed else ""
    return f"""
<section class="page-break">
  <h2>Part 2 — Key matters</h2>
  <table class="data-table">
    <tr><th>Landlord</th><td>{{{{landlord_full_name}}}}{{{{#if landlord_2_full_name}}}} and {{{{landlord_2_full_name}}}}{{{{/if}}}}</td></tr>
    <tr><th>Landlord contact address</th><td>{{{{landlord_address}}}}</td></tr>
    <tr><th>Telephone / email</th><td>{{{{landlord_phone}}}} / {{{{landlord_email}}}}</td></tr>
    <tr><th>Dwelling</th><td>{{{{property_address}}}}</td></tr>
  <tr><th>Occupation date</th><td>{{{{format_date tenancy_start_date}}}}</td></tr>
    {fixed_rows}
    <tr><th>Rent</th><td>{{{{currency rent_amount}}}} per {{{{rent_period}}}}</td></tr>
    <tr><th>First payment</th><td>{{{{currency first_payment}}}} on {{{{format_date first_payment_date}}}}</td></tr>
    <tr><th>Further payments</th><td>{{{{rent_due_day}}}} of each {{{{rent_period}}}}</td></tr>
    <tr><th>Deposit</th><td>{{{{currency deposit_amount}}}} — {{{{deposit_scheme_name}}}}</td></tr>
    <tr><th>Rent Smart Wales</th><td>{{{{rent_smart_wales_number}}}}</td></tr>
    {{{{#if occupation_exclusion_applies}}}}
    <tr><th>Occupation exclusion</th><td>From {{{{format_date occupation_exclusion_start_date}}}} to {{{{format_date occupation_exclusion_end_date}}}}</td></tr>
    {{{{/if}}}}
  </table>
  {tenant_rows("Contract-holder(s)")}
</section>"""


def build_wales(source_name: str, target_name: str, fixed: bool) -> None:
    doc = Document(SOURCES / source_name)
    if fixed:
        explanatory, _ = render_range(doc, 13, 39)
        terms, _ = render_range(doc, 98, len(doc.paragraphs), heading_counter=1, fixed_wales_terms=True)
        title = "Fixed Term Standard Occupation Contract"
        subtitle = "Wales — term of less than seven years"
        source = "Welsh Government fixed-term model written statement (June 2026) used as the clause-parity specification."
    else:
        explanatory, _ = render_range(doc, 18, 70)
        terms, _ = render_range(doc, 139, len(doc.paragraphs))
        title = "Periodic Standard Occupation Contract"
        subtitle = "Wales"
        source = "Welsh Government periodic model written statement (May 2026) used as the clause-parity specification."

    body = f"""
<section class="page-break"><h2>Part 1 — Explanatory information</h2>{explanatory}</section>
{wales_key_matters(fixed)}
<section class="page-break"><h2>Part 3 — Fundamental and supplementary terms</h2>{terms}</section>
{signature_block("Contract-holder")}"""
    write_template(WALES_TEMPLATES / target_name, shell(title, subtitle, body, source))


def scotland_key_matters() -> str:
    return f"""
<section class="page-break">
  <h2>Key tenancy details</h2>
  {tenant_rows("1. Tenant(s)")}
  {{{{#if agent_name}}}}<h2>2. Letting agent</h2>
  <table class="data-table">
    <tr><th>Name and address</th><td>{{{{agent_name}}}}<br>{{{{agent_address}}}}</td></tr>
    <tr><th>Contact</th><td>{{{{agent_phone}}}} / {{{{agent_email}}}}</td></tr>
    <tr><th>Registration</th><td>{{{{agent_registration_number}}}}</td></tr>
    <tr><th>Services</th><td>{{{{agent_services}}}}</td></tr>
    <tr><th>First contact for</th><td>{{{{agent_contact_matters}}}}</td></tr>
  </table>{{{{/if}}}}
  <h2>3. Landlord</h2>
  <table class="data-table">
    <tr><th>Name and address</th><td>{{{{landlord_full_name}}}}<br>{{{{landlord_address}}}}</td></tr>
    <tr><th>Contact</th><td>{{{{landlord_phone}}}} / {{{{landlord_email}}}}</td></tr>
    <tr><th>Registration</th><td>{{{{landlord_registration_number}}}}</td></tr>
    {{{{#if landlord_2_full_name}}}}<tr><th>Joint landlord</th><td>{{{{landlord_2_full_name}}}}<br>{{{{landlord_2_address}}}}<br>{{{{landlord_2_phone}}}} / {{{{landlord_2_email}}}}<br>Registration: {{{{landlord_2_registration_number}}}}</td></tr>{{{{/if}}}}
  </table>
  <h2>4. Communication</h2><p>Formal communications will be made in writing by: <strong>{{{{communication_method}}}}</strong>.</p>
  <h2>5. Details of the let property</h2>
  <table class="data-table">
    <tr><th>Address</th><td>{{{{property_address}}}}</td></tr>
    <tr><th>Type / furnishing</th><td>{{{{property_type}}}} / {{{{furnished_status}}}}</td></tr>
    <tr><th>Included areas</th><td>{{{{included_areas}}}}</td></tr>
    <tr><th>Shared areas</th><td>{{{{shared_areas}}}}</td></tr>
    <tr><th>Excluded areas</th><td>{{{{excluded_areas}}}}</td></tr>
    <tr><th>Rent pressure zone</th><td>{{{{#if in_rent_pressure_zone}}}}Yes{{{{else}}}}No{{{{/if}}}}</td></tr>
    <tr><th>House in Multiple Occupation</th><td>{{{{#if is_hmo}}}}Yes{{{{else}}}}No{{{{/if}}}}</td></tr>
  </table>
  <h2>6. Start date of the tenancy</h2><p>The private residential tenancy will start on <strong>{{{{format_date tenancy_start_date}}}}</strong>.</p>
</section>"""


def build_scotland() -> None:
    doc = Document(SOURCES / "scotland-model-prt-2024-04.docx")
    glossary, _ = render_range(doc, 61, 92)
    clauses_9_10, _ = render_range(doc, 253, 265, heading_counter=9)
    clauses_12_36, _ = render_range(
        doc,
        288,
        603,
        heading_counter=12,
        replacements={
            544: '<p class="legal-paragraph">The Tenant agrees that the signed Inventory and Record of Condition, '
            '{{#if_eq inventory_delivery_method "later"}}which will be supplied to the Tenant no later than the start date of the tenancy'
            '{{else}}attached as Schedule 1 to this Agreement{{/if_eq}}, '
            'is a full and accurate record of the contents and condition of the Let Property at the start date of the tenancy. '
            'The Tenant has 7 days from the start date to notify the Landlord in writing of discrepancies; otherwise the Tenant '
            'will be deemed satisfied with it.</p>',
            556: '<p class="legal-paragraph">The Tenant undertakes to place the accounts for '
            '<strong>{{tenant_utility_accounts}}</strong> in their name with the relevant supplier and to pay all sums due for '
            'those supplies during the tenancy.</p>',
        },
    )
    body = f"""
<section class="page-break"><h2>Glossary and interpretation</h2>{glossary}</section>
{scotland_key_matters()}
<section class="page-break">
  <h2>Model private residential tenancy terms</h2>
  <h3 class="term-heading"><span class="term-number">7.</span>Occupation and use of the let property</h3>
  <p class="legal-paragraph">The Tenant agrees to continue to occupy the Let Property as his or her home and must obtain the Landlord’s written permission before carrying out any trade, business or profession there.</p>
  <h3 class="term-heading"><span class="term-number">8.</span>Rent</h3>
  <p class="legal-paragraph">The rent is <strong>{{{{currency rent_amount}}}}</strong> per {{{{rent_period}}}}, payable in {{{{rent_payment_timing}}}}.</p>
  <p class="legal-paragraph">The first payment of {{{{currency first_payment}}}} will be paid on {{{{format_date first_payment_date}}}} and covers {{{{format_date first_payment_period_from}}}} to {{{{format_date first_payment_period_to}}}}. Thereafter {{{{currency rent_amount}}}} must be received on {{{{rent_due_day}}}} of each {{{{rent_period}}}}. Payment method: {{{{payment_method}}}}.</p>
  <p class="legal-paragraph">Services included in the rent: {{{{rent_includes}}}}.</p>
  {clauses_9_10}
  <h3 class="term-heading"><span class="term-number">11.</span>Deposit</h3>
  <p class="legal-paragraph">The Landlord must lodge any deposit received with a tenancy deposit scheme within 30 working days of the start date of the tenancy. Each later instalment must be lodged within 30 working days of payment.</p>
  <p class="legal-paragraph">A deposit of {{{{currency deposit_amount}}}} will be paid. It cannot exceed two months’ rent. The scheme administrator is <strong>{{{{deposit_scheme_name}}}}</strong>{{{{#if deposit_scheme_contact_details}}}}, contact: {{{{deposit_scheme_contact_details}}}}{{{{/if}}}}.</p>
  <p class="legal-paragraph">At the end of the tenancy the Landlord should request release through the scheme. The scheme’s dispute-resolution process applies where the amount is disputed; any liability above the protected deposit remains recoverable.</p>
  {clauses_12_36}
  <h3 class="term-heading"><span class="term-number">37.</span>Additional tenancy terms</h3>
  <p class="omitted">No additional discretionary terms are included in this standard agreement.</p>
  <h3 class="term-heading"><span class="term-number">38.</span>The guarantor</h3>
  <p class="omitted">No guarantor provision is included in this standard agreement.</p>
  <h3 class="term-heading"><span class="term-number">39.</span>Declarations</h3>
  <p class="legal-paragraph">By signing and taking entry to the Let Property, the Tenant confirms that they made full and true disclosure of information sought in connection with the tenancy, did not knowingly or carelessly make a false or misleading statement affecting the decision to grant it, and read and understood this Agreement and the accompanying statutory supporting notes.</p>
</section>
{signature_block("Tenant")}"""
    source = "Scottish Government Model Private Residential Tenancy Agreement (April 2024) used as the clause-parity specification. The official statutory supporting notes are supplied separately."
    write_template(
        SCOTLAND_TEMPLATES / "prt_agreement.hbs",
        shell("Private Residential Tenancy Agreement", "Scotland — open-ended tenancy", body, source),
    )


if __name__ == "__main__":
    build_wales(
        "wales-periodic-standard-2026-05.docx",
        "standard_occupation_contract.hbs",
        fixed=False,
    )
    build_wales(
        "wales-fixed-standard-2026-06.docx",
        "fixed_term_standard_occupation_contract.hbs",
        fixed=True,
    )
    build_scotland()
    print("Branded Wales periodic/fixed and Scotland PRT templates rebuilt.")
