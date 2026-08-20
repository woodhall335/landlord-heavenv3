from pathlib import Path
from shutil import copytree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, NumberObject, TextStringObject


SOURCE = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v11")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v12")

if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
copytree(SOURCE, ROOT)

service = ROOT / "01_SERVE_ON_20_AUGUST_2026"


def create_cover_note(path):
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    def add(text, bold=False, alignment=None, after=8):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(after)
        if alignment is not None:
            paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.bold = bold
        return paragraph

    add("Suraraj Pradhan", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add("26A Rhodes Place\nOldbrook\nMilton Keynes\nMK6 2LX\npradhansuraraj@gmail.com | 07747 817502", alignment=WD_ALIGN_PARAGRAPH.RIGHT, after=14)
    add("20 August 2026", after=14)
    add("To: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne", after=0)
    add("39 Upton Grove\nShenley Lodge\nMilton Keynes\nMK5 7GR", after=16)
    add("RE: RENTERS' RIGHTS ACT INFORMATION SHEET 2026", bold=True, after=14)
    add("Dear Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne,", after=10)
    add("Enclosed is the exact official Renters' Rights Act Information Sheet 2026.")
    add("I provide it for completeness and without admission as to whether it was previously supplied by GOA Property Solutions Ltd or any other former managing agent. It is enclosed so that you have the statutory information concerning changes to the tenancy under the Renters' Rights Act 2025.")
    add("This covering letter does not vary or replace the tenancy agreement, Form 3A notice or any other document served with it. Please retain the enclosed Information Sheet for your records.", after=16)
    add("Yours sincerely,", after=24)
    add("Signed: ________________________________", after=0)
    add("Suraraj Pradhan\nLandlord and issuer of this letter", after=0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Covering letter issued personally by Suraraj Pradhan"
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
    document.save(path)


create_cover_note(service / "07-renters-rights-act-information-sheet-covering-letter-for-signature.docx")


def replace_text(path, old, new):
    document = Document(path)
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            document.save(path)
            return
    raise SystemExit(f"Text not found in {path.name}")


replace_text(
    service / "02-service-instructions-and-checklist.docx",
    "The official Renters' Rights Act Information Sheet 2026 PDF is also enclosed. If it was not previously given to every tenant, give the exact PDF to each named tenant now as a separate document and keep proof. It does not form part of the Form 3A notice.",
    "The official Renters' Rights Act Information Sheet 2026 PDF and a landlord-issued covering letter are also enclosed. If the Information Sheet was not previously given to every tenant, Suraraj should sign the covering letter and give the exact PDF and covering letter to each named tenant now as separate documents. Keep proof. Neither document forms part of the Form 3A notice.",
)
replace_text(
    service / "04-service-record-for-deliverer-complete-immediately.docx",
    "(4) the exact official Renters' Rights Act Information Sheet 2026 PDF, if not already given to every tenant.",
    "(4) the exact official Renters' Rights Act Information Sheet 2026 PDF and its landlord-issued covering letter, if not already given to every tenant.",
)

n215 = ROOT / "02_COMPLETE_AFTER_SERVICE" / "01-n215-ground-1-service-certificate-editable-complete-after-service.pdf"
reader = PdfReader(str(n215))
writer = PdfWriter()
writer.clone_document_from_reader(reader)
writer.reattach_fields()
for page in writer.pages:
    for annotation in page.get("/Annots", []):
        widget = annotation.get_object()
        if widget.get("/T") == "Text1":
            widget[NameObject("/Ff")] = NumberObject(int(widget.get("/Ff", 0)) | 4096)
            widget[NameObject("/DA")] = TextStringObject("/ArialMT 8 Tf 0 g")
for page in writer.pages:
    writer.update_page_form_field_values(page, {
        "Text1": "Form 3A notice seeking possession (Ground 1), Ground 1 continuation sheet, formal GOA-revocation notice if served, and the Renters' Rights Act Information Sheet 2026 with its covering letter if served."
    }, auto_regenerate=True)
with open(n215, "wb") as output:
    writer.write(output)

print(ROOT)
