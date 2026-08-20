from pathlib import Path
from shutil import copy2, copytree

from docx import Document


SOURCE = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v9")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v10")
TITLE_SOURCE = Path(r"C:\Users\t_moh\Downloads\2026-04-02_summary_of_title_BM273247_GOV.UK.pdf")

if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
copytree(SOURCE, ROOT)

TITLE_DIR = ROOT / "04_EVIDENCE_BUNDLE" / "01_Standing_and_Title"
TITLE_DIR.mkdir()
copy2(TITLE_SOURCE, TITLE_DIR / "01-supplied-title-register-summary-bm273247-not-official-copy.pdf")


def replace_paragraphs(path, replacements, additions=None):
    document = Document(path)
    found = set()
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
                found.add(old)
    missing = set(replacements) - found
    if missing:
        raise SystemExit(f"Replacement text not found in {path.name}: {missing}")
    if additions:
        for text in additions:
            document.add_paragraph(text)
    document.save(path)


index = ROOT / "00_READ_FIRST_CASE_SUMMARY_AND_INDEX"
court = ROOT / "03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES"
evidence = ROOT / "04_EVIDENCE_BUNDLE"
witness = ROOT / "05_WITNESS_STATEMENT_FOR_SIGNATURE"

replace_paragraphs(
    index / "00-case-summary-and-merits-status.docx",
    {
        "Status: this is a prepared service pack, not a court-ready claim. Standing must be resolved by evidence and, because the written agreement identifies GOA as landlord, independent legal advice before issue is strongly recommended. Actual service proof, title evidence, a signed current witness statement and confirmed deposit position also remain essential.":
        "Status: this is a prepared service pack, not a court-ready claim. A supplied HM Land Registry title-register summary dated 2 April 2026 (title BM273247) records Suraraj as proprietor of the Property, but it expressly states it is not an official copy. An official copy remains required before issue. Standing must also be resolved by evidence and, because the written agreement identifies GOA as landlord, independent legal advice before issue is strongly recommended. Actual service proof, an official title register and a signed current witness statement remain essential."
    },
)

replace_paragraphs(
    index / "02-filing-status-and-next-steps.docx",
    {
        "2. Obtain the official title register and the documents showing GOA's role and the basis on which Suraraj is the proper claimant. The AST naming GOA creates a standing risk which a self-declaration does not cure.":
        "2. The supplied title-register summary dated 2 April 2026 (title BM273247) has been added to SP1. It records Suraraj as proprietor but is marked not an official copy. Obtain an official title register and the documents showing GOA's role and the basis on which Suraraj is the proper claimant. The AST naming GOA creates a standing risk which a self-declaration does not cure."
    },
)

replace_paragraphs(
    court / "01-court-issue-conditions-and-evidence-schedule.docx",
    {
        "Required Ground 1 evidence: official title register; evidence and independent advice confirming that Suraraj is entitled to sue despite the AST naming GOA; current tenancy agreement; signed Form 3A and continuation sheet; completed N215 and service record; signed current witness statement; current-address evidence; and evidence of genuine continuing intention to occupy as sole/main home.":
        "Required Ground 1 evidence: official title register; evidence and independent advice confirming that Suraraj is entitled to sue despite the AST naming GOA; current tenancy agreement; signed Form 3A and continuation sheet; completed N215 and service record; signed current witness statement; current-address evidence; and evidence of genuine continuing intention to occupy as sole/main home. The supplied 2 April 2026 title-register summary in SP1 is useful preliminary evidence only and must not be substituted for an official copy."
    },
)

replace_paragraphs(
    evidence / "00-exhibit-schedule-and-bundle-map.docx",
    {
        "SP1 - Official Land Registry title register/title deed and signed claimant's statement of standing and ownership. Status: title register requested; statement ready for signature. Independent standing advice remains required before issue because the AST identifies GOA as landlord.":
        "SP1 - Supplied HM Land Registry title-register summary dated 2 April 2026, title BM273247, and signed claimant's statement of standing and ownership. The summary records Suraraj Pradhan as proprietor of 39 Upton Grove but is marked not an official copy. Status: preliminary title evidence held; official copy still required before issue. Independent standing advice remains required because the AST identifies GOA as landlord."
    },
)

replace_paragraphs(
    witness / "00-claimants-statement-of-standing-and-ownership-for-signature.docx",
    {
        "2. I attach an official copy of the Land Registry title register for the Property at exhibit SP1.":
        "2. I attach the supplied HM Land Registry title-register summary dated 2 April 2026, title BM273247, at exhibit SP1. It records me as proprietor of the Property but states that it is not an official copy. I will provide an official copy before court issue."
    },
)

replace_paragraphs(
    witness / "01-ground-1-witness-statement-for-review-and-signature.docx",
    {
        "3. I am the registered freehold proprietor of the Property. I rely on the official title register and my signed statement of standing and ownership at SP1.":
        "3. The supplied HM Land Registry title-register summary dated 2 April 2026, title BM273247, records me as proprietor of the Property. It is marked not an official copy. I rely on it and my signed statement of standing and ownership at SP1, and will provide an official copy before court issue."
    },
)

(evidence / "00-evidence-index.txt").write_text(
    "EVIDENCE BUNDLE INDEX\n\n"
    "SP1 Standing and ownership: supplied HM Land Registry title-register summary dated 2 April 2026, title BM273247, and signed claimant's standing statement. The supplied summary is marked not an official copy; obtain an official copy before court issue. Independent standing advice remains required before issue.\n"
    "SP2 Current address: council-tax and utility evidence.\n"
    "SP3 Tenancy: written agreement dated 25 November 2025. It records the original AST terms; it became assured periodic on 1 May 2026.\n"
    "SP4 Management: GOA management agreement.\n"
    "SP5 GOA termination: termination notice to GOA.\n"
    "SP6 Service: signed Form 3A, continuation sheet, formal GOA-revocation notice if served, N215, service record and delivery proof.\n"
    "SP7 Ground 1 evidence: signed statement of intent, current witness statement and school-related evidence if available.\n"
    "SP8 Compliance: gas safety certificate, EICR and EPC. These are background/compliance evidence, not primary Ground 1 proof.\n"
    "SP9 Historic tenant-notification correspondence: context only, not proof of a formal notice served on all tenants.\n\n"
    "Do not issue until the official title register, standing evidence/advice, actual service proof, signed witness evidence and Renters' Rights Act information-sheet compliance check have been completed.\n",
    encoding="utf-8",
)

print(ROOT)
