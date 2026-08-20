# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2, copytree
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from pypdf import PdfReader, PdfWriter

ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-19-august-2026")
SRC = Path(r"C:\Users\t_moh\Documents\GitHub\landlord-heavenv3\tmp\suraraj-case-source")
FORMS = Path(r"C:\Users\t_moh\Documents\GitHub\landlord-heavenv3\tmp\suraraj-authoring-sources")
WILL_TEMPLATE = Path(r"C:\Users\t_moh\Downloads\will-court-pack\00_READ_FIRST_CASE_SUMMARY_AND_INDEX\04-cover-letter-to-william-and-leanne.docx")

for folder in [
    ROOT / '00_READ_FIRST_CASE_SUMMARY_AND_INDEX',
    ROOT / '01_SERVE_ON_19_AUGUST_2026',
    ROOT / '02_COMPLETE_AFTER_SERVICE',
    ROOT / '03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES',
    ROOT / '04_EVIDENCE_BUNDLE',
    ROOT / '05_WITNESS_STATEMENT_FOR_SIGNATURE',
]: folder.mkdir(parents=True, exist_ok=True)

def clear_paragraphs(doc):
    body = doc._element.body
    for p in list(body.findall(qn('w:p'))): body.remove(p)

def set_margins(doc):
    sec=doc.sections[0]
    sec.top_margin=Pt(54); sec.bottom_margin=Pt(54); sec.left_margin=Pt(58); sec.right_margin=Pt(58)

def heading(doc, text, level=1):
    p=doc.add_paragraph(style='Heading %d'%level)
    p.add_run(text)
    return p

def para(doc, text='', bold_prefix=None):
    p=doc.add_paragraph(style='Normal')
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p

def bullets(doc, items):
    for item in items: doc.add_paragraph(item, style='List Bullet')

def footer(doc):
    p=doc.sections[0].footer.paragraphs[0]
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.text='Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026'
    for r in p.runs: r.font.size=Pt(8)

def base_doc():
    copy2(WILL_TEMPLATE, ROOT/'tmp-template.docx')
    doc=Document(ROOT/'tmp-template.docx')
    clear_paragraphs(doc); set_margins(doc); footer(doc)
    return doc

# cover letter
doc=base_doc()
para(doc, 'Ground 1 possession pack - 39 Upton Grove')
para(doc, 'Dear Suraraj,')
para(doc, 'This pack has been prepared for a new Section 8 possession route relying on Ground 1 only. It does not rely on the previous notice and it does not allege Ground 8 rent arrears.')
para(doc, 'The Form 3A is dated for service on 19 August 2026 and gives more than four months’ notice. It names Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne. The landlord must sign the notice before service and use the service instructions in this pack.')
para(doc, 'The court forms are included only for the stage after the notice expires. Do not issue a claim until the evidence checklist has been completed, particularly the title/standing confirmation, deposit confirmation, signed witness statement and service evidence.')
para(doc, 'The current Ground 1 statement has been updated to remove the allegation that GOA granted the tenancy without consent. The signed management agreement contains express letting authority, so that allegation should not be used in court documents.')
para(doc, 'Please review every document carefully before signing or serving. The court, not Landlord Heaven, decides whether to make a possession order.')
para(doc, 'Yours sincerely')
para(doc, 'Tariq Mohammed\nLandlord Heaven')
doc.save(ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'01-cover-letter-to-suraraj.docx')

# status & next steps
doc=base_doc(); heading(doc,'Filing status and next steps')
heading(doc,'Current status')
para(doc,'A new private-sector Form 3A has been prepared on Ground 1 only. The earliest date for court proceedings is 19 December 2026, subject to the notice being signed and served on 19 August 2026.')
heading(doc,'Before service')
bullets(doc,[
'Check the tenant names, property address and landlord contact details against the tenancy agreement.',
'Suraraj Pradhan must sign and date the Form 3A on the actual date it is served.',
'Use no previous notice as the legal basis for this claim.',
'Keep a complete, identical copy of every page served.'
])
heading(doc,'Before court issue')
bullets(doc,[
'Obtain official title evidence and retain the solicitor confirmation that Suraraj is entitled to act as landlord.',
'Obtain written confirmation that no tenancy deposit was paid; if a deposit was paid, obtain scheme and prescribed-information evidence or legal advice before issue.',
'Use the updated Ground 1 witness statement only after Suraraj reviews, signs and dates it.',
'Complete the N215 with actual service facts; do not guess dates, times or methods.',
'Review any tenant response, disrepair allegation, vulnerability, equality issue, counterclaim or change in the landlord’s intention before issuing a claim.'
])
doc.save(ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'02-filing-status-and-next-steps.docx')

# service instructions
doc=base_doc(); para(doc,'Form 3A Ground 1 service checklist')
para(doc,'SERVICE ON 19 AUGUST 2026 ONLY')
heading(doc,'Serve')
bullets(doc,[
'Serve a complete copy of the signed Form 3A and its Ground 1 continuation sheet on all three tenants at 39 Upton Grove: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
'Use hand delivery to the property and send the same complete pack by a tracked postal method on the same day.',
'Do not enter the property or seek to force a discussion. Do not change locks or take any step to exclude the tenants.'
])
heading(doc,'Record immediately')
bullets(doc,[
'Exact date and time of hand delivery.',
'Name of the person who delivered the documents and a signed contemporaneous note.',
'Photograph of the addressed envelope at the property, if safely possible, and the postal receipt/tracking number.',
'Copies of every page served and the signed Form 3A.'
])
heading(doc,'After service')
bullets(doc,[
'Complete the current N215 using actual service facts only.',
'Do not apply to court before 19 December 2026.',
'If the tenants remain after the notice date, complete the court evidence checklist before considering N5/N119 issue.'
])
doc.save(ROOT/'01_SERVE_ON_19_AUGUST_2026'/'02-service-instructions-and-checklist.docx')

# court issue checklist
doc=base_doc(); heading(doc,'Court issue conditions and evidence schedule')
para(doc,'This schedule is for the post-notice court stage. It is not a claim form and it does not authorise issue.')
heading(doc,'Core Ground 1 evidence')
bullets(doc,[
'Official title evidence for 39 Upton Grove.',
'Solicitor confirmation that Suraraj is the current landlord entitled to seek possession.',
'Signed Ground 1 witness statement and current-address evidence.',
'Proof of genuine intention to occupy as sole or principal home.',
'Copy of the signed Form 3A and completed N215/service evidence.',
'Current tenancy agreement naming all three tenants.'
])
heading(doc,'Compliance and risk checks')
bullets(doc,[
'Written confirmation that no tenancy deposit was taken; otherwise deposit scheme, prescribed information and remedial advice.',
'Gas safety, EPC and electrical evidence, together with evidence of any required delivery.',
'Any tenant correspondence, defence, payment record, complaint, repair issue or vulnerability disclosure received after service.',
'Confirmation that the landlord still intends to occupy the property as their only or principal home.'
])
heading(doc,'Court forms')
para(doc,'The current official N5, N119 and N215 forms are included as court-stage source forms. They must be completed from the final evidence and signed at the relevant stage; they are not issued with this notice.')
doc.save(ROOT/'03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'/'01-court-issue-conditions-and-evidence-schedule.docx')

# witness statement
doc=base_doc(); heading(doc,'Witness statement of Suraraj Pradhan')
para(doc,'IN THE COUNTY COURT AT [TO BE COMPLETED IF CLAIM ISSUED]')
para(doc,'Claim No: [TO BE COMPLETED]')
para(doc,'Between: Suraraj Pradhan - Claimant\nand\nShellyann Roberts-Henry, Mignal Small and Taylor Goulbourne - Defendants')
heading(doc,'Statement')
for n,t in enumerate([
'I am Suraraj Pradhan of 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX. I am the freehold owner of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property).',
'I seek possession under Ground 1 of Schedule 2 to the Housing Act 1988. I require the Property as my only or principal home.',
'I currently live at 26A Rhodes Place. I rely on my council tax and utility evidence as confirmation of that address.',
'My son turns four in December 2026. I intend to move to the Property to support his school application; the intended school is approximately 200 metres from the Property and the Property is within its catchment area.',
'I intend to occupy the Property permanently as my sole and main residence on obtaining possession. I do not intend to use it as a second or holiday home.',
'The tenancy agreement dated 25 November 2025 names the three defendants. I understand that the tenancy is now periodic. I served a new Form 3A notice dated 19 August 2026 and rely only on Ground 1 in that notice.',
'I do not rely on the previous notice or on any allegation that the tenancy was granted without my consent. I understand that the court will determine the claim.'
],1): para(doc,f'{n}. {t}')
heading(doc,'Statement of truth')
para(doc,'I believe that the facts stated in this witness statement are true.')
para(doc,'Signed: ________________________________\nSuraraj Pradhan\n\nDated: _________________________________')
doc.save(ROOT/'05_WITNESS_STATEMENT_FOR_SIGNATURE'/'01-ground-1-witness-statement-for-review-and-signature.docx')

# create Form 3A overlay
form_source=FORMS/'Form_3A_0526.pdf'; overlay=ROOT/'tmp-form3a-overlay.pdf'; completed=ROOT/'01_SERVE_ON_19_AUGUST_2026'/'01-form-3a-ground-1-serve-19-08-2026.pdf'
c=canvas.Canvas(str(overlay),pagesize=A4)
def textbox(page_index,x,y,w,h,text,size=9,leading=None):
    # ReportLab y is bottom coordinate; page_index handles page advance
    while c.getPageNumber() < page_index+1: c.showPage()
    c.setFont('Helvetica',size)
    leading=leading or size*1.2
    words=text.split(); lines=[]; current=''
    for word in words:
        trial=(current+' '+word).strip()
        if stringWidth(trial,'Helvetica',size) <= w:
            current=trial
        else:
            lines.append(current); current=word
    if current: lines.append(current)
    top=y+h-size
    for line in lines:
        if top < y: break
        c.drawString(x,top,line); top-=leading
def tick(page_index,x,y):
    while c.getPageNumber() < page_index+1: c.showPage()
    c.setFont('Helvetica-Bold',11); c.drawString(x,y,'X')

# page 2 index 1
textbox(1,58,738,355,16,'Shellyann Roberts-Henry',9)
textbox(1,58,720,355,16,'Mignal Small',9)
textbox(1,58,702,355,16,'Taylor Goulbourne',9)
textbox(1,58,614,355,18,'39 Upton Grove',9)
textbox(1,58,573,355,18,'Shenley Lodge',9)
textbox(1,58,532,205,18,'Milton Keynes',9)
textbox(1,58,491,200,18,'',9)
textbox(1,58,450,155,18,'MK5 7GR',9)
while c.getPageNumber() < 2: c.showPage()
c.setFont('Helvetica',9)
c.drawString(96,283,'19')
c.drawString(134,283,'12')
c.drawString(172,283,'2026')
# page 3 tick ground 1 (third checkbox)
tick(2,59,387)
# legal wording p4
legal='''Ground 1\nThe current tenancy began at least 1 year before the relevant date and the landlord who is seeking possession requires the dwelling-house as the only or principal home of any of the following—\n(a) the landlord;\n(b) the landlord’s spouse or civil partner or a person with whom the landlord lives as if they were married or in a civil partnership;\n(c) the landlord’s parent, grandparent, sibling, child or grandchild; or\n(d) a child or grandchild of a person mentioned in paragraph (b).\nA relationship of the half-blood is to be treated as a relationship of the whole blood.\nIn the case of joint landlords seeking possession, references to “the landlord” in this ground are to be read as references to at least one of those joint landlords.\nWhen calculating whether the current tenancy began at least 1 year before the relevant date, both the day when the current tenancy began and the relevant date must be included in the calculation.'''
textbox(3,60,70,350,650,legal,8,10)
explanation='''Ground 1 - occupation by landlord\n\nThe landlord, Suraraj Pradhan, requires 39 Upton Grove as his only or principal home. He is the freehold owner and has confirmed that he is entitled to act as the current landlord following the termination of GOA Property Solutions Ltd’s management arrangement.\n\nMr Pradhan currently resides at 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX. He intends to occupy the Property permanently as his sole and main residence after possession is obtained. His son turns four in December 2026 and Mr Pradhan intends to move to support his son’s school application; the intended school is approximately 200 metres from the Property and the Property is within its catchment area.\n\nThe tenancy began on 25 November 2025. The earliest date for court proceedings in this notice is 19 December 2026, which is more than 12 months after the tenancy began.\n\nThe landlord relies on the tenancy agreement, title evidence, current-address evidence and his signed witness statement. This notice does not rely on the previous notice or on Ground 8.'''
textbox(4,60,65,350,670,explanation,8,10)
# contact p6
tick(5,59,638)
textbox(5,58,505,175,18,'19 August 2026',9)
textbox(5,58,431,350,22,'Suraraj Pradhan',9)
textbox(5,58,341,350,18,'26A Rhodes Place',9)
textbox(5,58,300,350,18,'Oldbrook',9)
textbox(5,58,259,205,18,'Milton Keynes',9)
textbox(5,58,217,205,18,'',9)
textbox(5,58,177,155,18,'MK6 2LX',9)
textbox(5,58,118,170,18,'07747 817502',9)
textbox(5,58,53,350,24,'pradhansuraraj@gmail.com',9)
# advance to ten pages
while c.getPageNumber() < 10: c.showPage()
c.save()
base=PdfReader(str(form_source)); over=PdfReader(str(overlay)); writer=PdfWriter()
for i,page in enumerate(base.pages):
    if i < len(over.pages): page.merge_page(over.pages[i])
    writer.add_page(page)
with open(completed,'wb') as f: writer.write(f)

# official forms for court stage and after service
copy2(FORMS/'N5_0526.pdf', ROOT/'03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'/'02-current-official-n5-england-05-26.pdf')
copy2(FORMS/'N119_0526.pdf', ROOT/'03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'/'03-current-official-n119-england-05-26.pdf')
copy2(FORMS/'N215_0626.pdf', ROOT/'02_COMPLETE_AFTER_SERVICE'/'01-current-n215-complete-after-service.pdf')

# evidence copying with concise index
manifest=[]
def add(src,sub,name):
    dst=ROOT/'04_EVIDENCE_BUNDLE'/sub/name; dst.parent.mkdir(parents=True,exist_ok=True); copy2(src,dst); manifest.append(f'{sub}/{name}')
add(SRC/'email-archive-2'/'Assured Shorthold Tenancy Agreement - 39 Upton Grove - Signed.pdf','01_Tenancy','01-ast-signed.pdf')
add(SRC/'email-archive-2'/'2026-04-01-Formal Termination Notice- GOA Properties.pdf','02_Management_and_Standing','01-termination-notice-to-goa.pdf')
add(Path(r'C:\Users\t_moh\Downloads\39 Upton Grove - Management Agreement - SGA-Signed.pdf'),'02_Management_and_Standing','02-management-agreement.pdf')
add(SRC/'email-archive-1'/'26A Rhodes Place-Council Tax - 2026-2027.pdf','03_Ground_1','01-current-address-council-tax.pdf')
add(SRC/'email-archive-1'/'Utility bill-2026-06-23-to-2026-07-22.pdf','03_Ground_1','02-current-address-utility.pdf')
add(SRC/'email-archive-2'/'Gas Safety Certificate -2-Valid from -17-Apr-2026 to 17-Apr-2027.pdf.pdf','04_Compliance','01-gas-safety-certificate.pdf')
add(SRC/'email-archive-2'/'EICR-68710-SurarajPradhan-6871000001078-968774687.pdf','04_Compliance','02-eicr.pdf')
add(SRC/'email-archive-2'/'Energy performance certificate (EPC) – Find an energy certificate – GOV.UK.pdf','04_Compliance','03-epc.pdf')
(Path(ROOT/'04_EVIDENCE_BUNDLE'/'00-evidence-index.txt')).write_text('\n'.join(manifest)+'\n\nDo not file until the court-stage checklist is complete.\n',encoding='utf-8')
(ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'00-IMPORTANT-READ-ME.txt').write_text('''GROUND 1 SERVICE PACK\n\nThis pack is prepared for a new Ground 1 Form 3A service on 19 August 2026.\n\nThe notice is not valid unless signed and served on that date. It is Ground 1 only. Ground 8 is deliberately excluded because the supplied records show tenant payments to GOA and Black & White, not proved tenant arrears owed to Suraraj.\n\nDo not issue N5/N119 until the court-stage evidence schedule is complete and the notice has expired.\n''',encoding='utf-8')
(ROOT/'01_SERVE_ON_19_AUGUST_2026'/'03-DO-NOT-SERVE-UNLESS-CHECKLIST-COMPLETE.txt').write_text('''Before serving: confirm Form 3A is signed by Suraraj Pradhan, all pages are included, the date is 19 August 2026, and a full identical copy is retained. Serve every page on all three tenants.\n''',encoding='utf-8')

# Remove temporary template copy
(ROOT/'tmp-template.docx').unlink(missing_ok=True)
print(ROOT)

