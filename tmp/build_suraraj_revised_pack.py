# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2, copytree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, NumberObject, TextStringObject
from reportlab.lib.colors import white
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

SOURCE = Path(r'C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-19-august-2026')
ROOT = Path(r'C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v3')
FORMS = Path(r'C:\Users\t_moh\Documents\GitHub\landlord-heavenv3\tmp\suraraj-authoring-sources')
CASE_SOURCE = Path(r'C:\Users\t_moh\Documents\GitHub\landlord-heavenv3\tmp\suraraj-case-source')

if ROOT.exists():
    raise SystemExit(f'Remove or rename existing destination before build: {ROOT}')
copytree(SOURCE, ROOT)

old_service = ROOT / '01_SERVE_ON_19_AUGUST_2026'
service = ROOT / '01_SERVE_ON_20_AUGUST_2026'
old_service.rename(service)

court = ROOT / '03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'
evidence = ROOT / '04_EVIDENCE_BUNDLE'
index = ROOT / '00_READ_FIRST_CASE_SUMMARY_AND_INDEX'
template = index / '01-cover-letter-to-suraraj.docx'


def make_doc(footer='Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026'):
    document = Document(template)
    body = document._element.body
    for paragraph in list(body.findall(qn('w:p'))):
        body.remove(paragraph)
    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.text = footer
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)
    return document


def add_lines(document, lines):
    for text in lines:
        document.add_paragraph(text)


def save_document(filename, lines):
    document = make_doc()
    add_lines(document, lines)
    document.save(filename)


def replace_document(filename, lines):
    save_document(filename, lines)


def make_font_ready_reader(source, temporary, draw_static=None):
    reader = PdfReader(str(source))
    box = reader.pages[0].mediabox
    overlay = canvas.Canvas(str(temporary), pagesize=(float(box.width), float(box.height)))
    overlay.setFillColor(white)
    overlay.setFont('Helvetica', 1)
    for page_number in range(len(reader.pages)):
        if draw_static:
            draw_static(overlay, page_number)
        overlay.drawString(0, 0, ' ')
        overlay.showPage()
    overlay.save()
    overlay_reader = PdfReader(str(temporary))
    merged = PdfWriter()
    for page_number, page in enumerate(reader.pages):
        page.merge_page(overlay_reader.pages[page_number])
        merged.add_page(page)
    prepared = temporary.with_name(temporary.stem + '-prepared.pdf')
    with open(prepared, 'wb') as output:
        merged.write(output)
    temporary.unlink(missing_ok=True)
    return PdfReader(str(prepared)), prepared


def fill_official_form(source, destination, values, font_ready=False, draw_static=None, multiline_fields=None, field_styles=None):
    temporary = None
    if font_ready:
        reader, temporary = make_font_ready_reader(source, ROOT / 'tmp-font-resource.pdf', draw_static)
    else:
        reader = PdfReader(str(source))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    writer.reattach_fields()
    if multiline_fields:
        for page in writer.pages:
            for annotation in page.get('/Annots', []):
                widget = annotation.get_object()
                if widget.get('/T') in multiline_fields:
                    flags = int(widget.get('/Ff', 0)) | 4096
                    widget[NameObject('/Ff')] = NumberObject(flags)
                if field_styles and widget.get('/T') in field_styles:
                    widget[NameObject('/DA')] = TextStringObject(field_styles[widget.get('/T')])
    for page in writer.pages:
        writer.update_page_form_field_values(page, values, auto_regenerate=True)
    with open(destination, 'wb') as output:
        writer.write(output)
    if temporary:
        temporary.unlink(missing_ok=True)


tenant_names = 'Shellyann Roberts-Henry\nMignal Small\nTaylor Goulbourne'
property_address = '39 Upton Grove\nShenley Lodge\nMilton Keynes\nMK5 7GR'
landlord_address = '26A Rhodes Place\nOldbrook\nMilton Keynes\nMK6 2LX'
ground_1_legal = (
    'Ground 1\n'
    'The current tenancy began at least 1 year before the relevant date and the landlord who is seeking possession '
    'requires the dwelling-house as the only or principal home of any of the following: (a) the landlord; (b) the '
    "landlord's spouse or civil partner or a person with whom the landlord lives as if they were married or in a civil "
    'partnership; (c) the landlord\'s parent, grandparent, sibling, child or grandchild; or (d) a child or grandchild '
    'of a person mentioned in paragraph (b). A relationship of the half-blood is to be treated as a relationship of '
    'the whole blood. In the case of joint landlords seeking possession, references to the landlord are to be read as '
    'references to at least one of those joint landlords. When calculating whether the current tenancy began at least '
    '1 year before the relevant date, both the day when the current tenancy began and the relevant date must be included.'
)
ground_1_reasons = (
    'Suraraj Pradhan requires the Property as his only or principal home. He currently lives at 26A Rhodes Place, '
    'Oldbrook, Milton Keynes, MK6 2LX. His son turns four in December 2026 and he intends to move to the Property '
    'to support the school application; the proposed school is about 200 metres away and the Property is in its catchment area. '
    'He intends to occupy the Property permanently as his sole and main home. The current tenancy began on 25 November 2025. '
    'The earliest date for proceedings in this notice is 20 December 2026, after the initial 12-month period. The claimant '
    'relies on title evidence, current-address evidence and his signed witness statement. He does not rely on Ground 8 or '
    'on any assertion that GOA granted the tenancy without consent. No previous notice is relied upon.'
)

form3a = service / '01-form-3a-ground-1-serve-20-08-2026-editable.pdf'
def draw_form3a_static(page, page_number):
    page.setFillColorRGB(0, 0, 0)
    page.setFont('Helvetica-Bold', 10)
    if page_number == 2:
        page.drawString(60, 388, 'X')
    if page_number == 5:
        page.drawString(60, 638, 'X')
fill_official_form(FORMS / 'Form_3A_0526.pdf', form3a, {
    'Text Field 132': tenant_names,
    'Text Field 116': '39 Upton Grove',
    'Text Field 115': 'Shenley Lodge',
    'Text Field 114': 'Milton Keynes',
    'Text Field 112': 'MK5 7GR',
    'Text Field 117': '20122026',
    'Check Box 71': '/Yes',
    'Text Field 133': 'Ground 1. The full statutory wording is set out in the attached continuation sheet, which forms part of this notice.',
    'Text Field 134': ground_1_reasons,
    'Check Box 42': '/Landlord',
    'Text Field 130': '20082026',
    'Text Field 135': 'Suraraj Pradhan',
    'Text Field 103': '26A Rhodes Place',
    'Text Field 102': 'Oldbrook',
    'Text Field 101': 'Milton Keynes',
    'Text Field 99': 'MK6 2LX',
    'Text Field 120': '07747 817502',
    'Text Field 104': 'pradhansuraraj@gmail.com',
}, font_ready=True, draw_static=draw_form3a_static)

continuation = service / '02-form-3a-continuation-sheet-ground-1-serve-with-notice.pdf'
styles = getSampleStyleSheet()
story = [
    Paragraph('Continuation Sheet to Form 3A - Ground 1', styles['Title']),
    Paragraph('Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR', styles['Normal']),
    Paragraph('Landlord: Suraraj Pradhan. Tenants: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.', styles['Normal']),
    Paragraph('Notice date: 20 August 2026. This continuation sheet forms part of the Form 3A notice and must be served with every copy of it.', styles['Normal']),
    Spacer(1, 5 * mm),
    Paragraph('Question 4.2 - Full legal wording: Ground 1', styles['Heading2']),
    Paragraph(ground_1_legal.replace('\n', '<br/>'), styles['Normal']),
    Spacer(1, 5 * mm),
    Paragraph('Question 4.3 - Why Ground 1 is used', styles['Heading2']),
    Paragraph(ground_1_reasons, styles['Normal']),
]
SimpleDocTemplate(str(continuation), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm).build(story)

n5 = court / '02-draft-n5-ground-1-pre-issue-editable.pdf'
def draw_n5_static(page, page_number):
    if page_number != 0:
        return
    page.setFillColorRGB(0, 0, 0)
    page.setFont('Helvetica-Bold', 10)
    page.drawString(32, 315, 'X')
    page.drawString(32, 255, 'X')
fill_official_form(FORMS / 'N5_0526.pdf', n5, {
    'Text Field 154': tenant_names,
    'Text Field 153': property_address,
    'Text Field 155': property_address,
    'Check Box 45': '/Yes',
    'Check Box 47': '/Yes',
    'Text Field 114': '26A Rhodes Place',
    'Text Field 113': 'Oldbrook',
    'Text Field 112': 'Milton Keynes',
    'Text13': 'MK6 2LX',
    'Text Field 172': 'pradhansuraraj@gmail.com',
    'Text Field 110': '07747 817502',
}, font_ready=True, draw_static=draw_n5_static, multiline_fields=['Text Field 154'], field_styles={'Text Field 154': '/ArialMT 7 Tf 0 g'})

n119 = court / '03-draft-n119-ground-1-pre-issue-editable.pdf'
def draw_n119_static(page, page_number):
    page.setFillColorRGB(0, 0, 0)
    page.setFont('Helvetica-Bold', 10)
    if page_number == 1:
        page.drawString(60, 470, 'X')
    if page_number == 3:
        page.drawString(60, 453, 'X')
fill_official_form(FORMS / 'N119_0526.pdf', n119, {
    'Text Field 3': 'Suraraj Pradhan',
    'Text Field 4': tenant_names,
    'Text Field 100': '39 Upton Grove',
    'Text Field 99': 'Shenley Lodge',
    'Text Field 98': 'Milton Keynes',
    'Text Field 96': 'MK5 7GR',
    'Text Field 108': tenant_names,
    'Text Field 109': 'Assured shorthold tenancy (written AST dated 25 November 2025)',
    'Text Field 102': '25112025',
    'Text Field 110': '1,700',
    'Check Box 27': '/Yes',
    'Text Field 113': (
        'The claimant seeks possession under Ground 1 of Schedule 2 to the Housing Act 1988. The claimant requires '
        'the Property as his only or principal home. The tenancy began on 25 November 2025. A new Form 3A dated '
        '20 August 2026 is prepared for service on each defendant. Before issue, this draft must be updated with actual '
        'service evidence. The earliest date for proceedings is 20 December 2026. The claimant relies on title evidence, '
        'current-address evidence, the tenancy agreement, service evidence and a signed witness statement. The claimant '
        'does not rely on Ground 8.'
    ),
    'Text Field 114': 'No rent arrears are claimed in this Ground 1-only claim.',
    'Check Box 51': '/Yes',
    'Text Field 115': 'Form 3A notice seeking possession',
}, font_ready=True, draw_static=draw_n119_static)

n215 = ROOT / '02_COMPLETE_AFTER_SERVICE' / '01-n215-ground-1-service-certificate-editable-complete-after-service.pdf'
fill_official_form(FORMS / 'N215_0626.pdf', n215, {
    'Text Field 3': 'Suraraj Pradhan',
    'Text Field 4': tenant_names,
    'Text Field 93': '20082026',
    'Text Field 94': '20082026',
    'Text1': 'Form 3A notice seeking possession (Ground 1), including all continuation pages.',
    'Text2': tenant_names,
    'Check Box4': '/Yes',
    'Text Field 102': '39 Upton Grove',
    'Text Field 101': 'Shenley Lodge',
    'Text Field 100': 'Milton Keynes',
    'Text16': 'MK5 7GR',
}, font_ready=True)

for old in (
    service / '01-form-3a-ground-1-serve-19-08-2026.pdf',
    court / '02-draft-n5-ground-1-pre-issue.pdf',
    court / '03-draft-n119-ground-1-pre-issue.pdf',
    ROOT / '02_COMPLETE_AFTER_SERVICE' / '01-current-n215-complete-after-service.pdf',
):
    old.unlink(missing_ok=True)

replace_document(index / '01-cover-letter-to-suraraj.docx', [
    'Ground 1 possession pack - 39 Upton Grove',
    'Dear Suraraj,',
    'This revised pack is prepared for a new Form 3A notice relying on Ground 1 only. It does not rely on a previous notice and it does not allege Ground 8 rent arrears.',
    'The editable official Form 3A is prepared for service on 20 August 2026. It names Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne, gives four months notice and identifies 20 December 2026 as the earliest date for proceedings. You must personally sign it on the actual day of service.',
    'The N5, N119 and N215 are editable official HMCTS forms. The N5 and N119 are pre-issue drafts only. The N215 is prepared for the planned service date but must be checked against what actually occurs, signed only after service and retained with the service record.',
    'This pack is not court-ready until the title/standing and deposit points are evidenced, the notice is validly served, and the signed witness statement and service evidence are complete. The court decides whether Ground 1 is proved.',
    'Yours sincerely',
    'Tariq Mohammed',
    'Landlord Heaven',
])

replace_document(index / '02-filing-status-and-next-steps.docx', [
    'Filing status and next steps',
    'Current status',
    'A new editable Form 3A has been prepared on Ground 1 only for service on 20 August 2026. The earliest date for court proceedings is 20 December 2026. The old notice is not relied upon.',
    'Before service on 20 August 2026',
    '1. Check the tenant names, property address and landlord contact details against the tenancy agreement.',
    '2. Obtain the official title register and evidence that GOA authority was revoked and notified to every tenant.',
    '3. Obtain written confirmation about any tenancy deposit. If a deposit was received, obtain scheme and prescribed-information evidence and take legal advice before issue.',
    '4. Suraraj Pradhan must sign and date the Form 3A on 20 August 2026. Serve a full identical copy on all three tenants and complete the service record immediately.',
    'After service',
    'Complete the N215 only with actual service facts. Keep the signed Form 3A, a copy of every page served, the service record, delivery proof and any tenant response.',
    'Before court issue',
    'Do not issue before 20 December 2026. Recheck the continuing intention to occupy, title/standing evidence, deposit position, witness statement, compliance material, service proof, tenant correspondence, repair issues, vulnerability/equality issues and any counterclaim.',
])

replace_document(index / '03-what-we-need-from-you-before-service.docx', [
    'What we need from you before Ground 1 service',
    'Dear Suraraj,',
    'Before serving the new Form 3A on 20 August 2026, please send:',
    '1. An official Land Registry title register/title deed for 39 Upton Grove, proving your ownership.',
    '2. The formal written notice sent to the tenants confirming that GOA Property Solutions Ltd authority had been revoked, together with evidence of service on Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    '3. Written confirmation that no tenancy deposit was taken. If a deposit was taken, send the amount, date, recipient, scheme details and prescribed information.',
    'The enclosed Form 3A is dated for service on 20 August 2026 only. Do not serve it on a later date. If these items cannot be checked before service, ask us to prepare a newly dated notice with a revised notice period.',
    'Yours sincerely',
    'Tariq Mohammed',
    'Landlord Heaven',
])

save_document(index / '00-case-summary-and-merits-status.docx', [
    'Case summary and merits status',
    'Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR.',
    'Claimant: Suraraj Pradhan. Defendants: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'Route: Ground 1, occupation by landlord or family. The case is deliberately not advanced on Ground 8 because the supplied material records tenant payments to GOA and Black & White and does not safely establish arrears owed by the tenants to Suraraj.',
    'Core case: Suraraj says he will occupy the Property as his sole or principal home. His son turns four in December 2026; the proposed school is about 200 metres away and the Property is within its catchment area. The tenancy began on 25 November 2025. A notice served on 20 August 2026 expires after the first 12 months and gives four months notice.',
    'Status: the legal route is arguable and the documentary case is organised, but it is not ready to issue. Ownership/standing, deposit position, service proof and a signed current witness statement remain essential conditions.',
])

save_document(index / '04-bundle-index-and-document-status.docx', [
    'Bundle index and document status',
    '1. Read first: case summary, cover letter, next steps, evidence request and merits/risk review.',
    '2. Serve on 20 August 2026: editable Form 3A, service instructions, service record and do-not-serve warning.',
    '3. Complete after service: editable N215 certificate of service. It is not to be signed until actual service has occurred.',
    '4. Court stage after 20 December 2026: editable N5 and N119 drafts, plus court-issue evidence schedule. They are not to be signed or filed before their final review.',
    '5. Evidence bundle: tenancy, GOA management/termination, current-address, compliance, signed Ground 1 source statement and historic tenant-notification correspondence retained for review.',
    '6. Witness evidence: updated Ground 1 witness statement for Suraraj to review, sign and date shortly before court issue.',
])

save_document(index / '05-merits-risks-and-response-plan.docx', [
    'Ground 1 merits, risks and response plan',
    'Merits: Ground 1 is mandatory if every statutory condition is proved. The proposed 20 August 2026 notice gives four months notice and its earliest proceedings date, 20 December 2026, is after the first 12 months of the tenancy that began on 25 November 2025.',
    'Strengths: all three named tenants are included; the current official Form 3A is used; the case gives a concrete sole/main-home reason; current-address evidence, tenancy evidence and compliance documents are held; Ground 8 is not pleaded on an unsafe payment theory.',
    'Material risks: the AST names GOA and the management agreement gave GOA letting authority, so Suraraj must retain title evidence and solicitor confirmation that he is the current landlord entitled to sue. Deposit evidence remains unresolved. The court will scrutinise whether the stated intention to occupy is genuine and still current. Poor service or a notice served late would undermine the claim.',
    'Anticipated response: the tenants may say that they paid rent, that GOA was the landlord, that notice was not properly served, that a deposit was not protected, or that the stated intention is not genuine. The Ground 1 case does not seek a rent judgment. The response is title/standing proof, service evidence, deposit evidence, the tenancy and management records, and Suraraj evidence of a genuine continuing intention to make the Property his sole/main home.',
    'Assessment: do not treat this as an 80% to 85% case at this stage. It becomes a credible mandatory Ground 1 claim only once the outstanding evidence is supplied and the notice is validly served. Without those items, it should not be issued.',
])

replace_document(service / '02-service-instructions-and-checklist.docx', [
    'Form 3A Ground 1 service instructions',
    'SERVICE ON 20 AUGUST 2026 ONLY',
    'Before leaving to serve, check that Suraraj has signed and dated the editable Form 3A, every continuation page is present, the tenant names are correct and the notice states 20 December 2026 as the earliest date for proceedings.',
    'Serve one complete identical copy addressed to all three tenants at 39 Upton Grove. Use an independent adult or professional process server to leave the notice at the Property. Do not enter, do not seek confrontation, do not change locks and do not attempt self-help eviction.',
    'Check the tenancy agreement for any contractual service clause before service. If it requires a specific method, follow it. Keep a copy of the notice exactly as served.',
    'Immediately record the actual time, exact manner of delivery, full name and contact details of the deliverer, the address used, a contemporaneous signed service record and any safe photograph of the addressed envelope or delivery. If a postal method is also used, retain the receipt and tracking evidence but do not describe a postal delivery as hand delivery.',
    'After service, complete the editable N215 and the service record using actual facts only. Suraraj must not sign the N215 statement of truth until satisfied it is accurate. Do not issue a court claim before 20 December 2026.',
])

save_document(service / '04-service-record-for-deliverer-complete-immediately.docx', [
    'Service record - complete immediately after service',
    'Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR.',
    'Documents served: complete Form 3A notice seeking possession (Ground 1), including all continuation pages.',
    'Persons named on the notice: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'Planned service date: 20 August 2026. Actual date: ____________________. Actual time: ____________________.',
    'Deliverer full name: ____________________. Address/contact number: ____________________.',
    'Exact service method and location: ________________________________________________________________.',
    'Was a copy also posted? Yes / No. If yes, method, receipt/tracking number and posting time: ________________________________________.',
    'Documents and proof retained: signed notice copy / all continuation pages / photograph if safe / postal proof if used / other: ____________________.',
    'I confirm this record is true and was made immediately after service. Deliverer signature: ____________________. Date: ____________________.',
])

replace_document(court / '01-court-issue-conditions-and-evidence-schedule.docx', [
    'Court issue conditions and evidence schedule',
    'This schedule is for the post-notice court stage. It does not authorise issue.',
    'Do not issue before 20 December 2026. Complete a final review first.',
    'Required Ground 1 evidence: official title evidence; solicitor confirmation of Suraraj standing as landlord; current tenancy agreement; signed Form 3A and completed N215/service record; signed current witness statement; current-address evidence; evidence of genuine continuing intention to occupy as sole/main home.',
    'Required risk checks: written confirmation of the deposit position; if a deposit was taken, scheme/protection and prescribed-information evidence or further legal advice; any tenant response, defence, repair complaint, vulnerability/equality issue, counterclaim or change in circumstances; evidence that the former-agent authority/notification point is properly addressed.',
    'Forms: N5 and N119 are editable pre-issue drafts. Court name, claim number, issue date, court fee, statement of truth/signature and service-specific fields must remain blank until the final filing review.',
])

replace_document(ROOT / '05_WITNESS_STATEMENT_FOR_SIGNATURE' / '01-ground-1-witness-statement-for-review-and-signature.docx', [
    'Witness statement of Suraraj Pradhan',
    'IN THE COUNTY COURT AT [TO BE COMPLETED IF CLAIM ISSUED]',
    'Claim No: [TO BE COMPLETED]',
    'Between: Suraraj Pradhan - Claimant',
    'and Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne - Defendants',
    '1. I am Suraraj Pradhan of 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX. I am the owner of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property). I will exhibit official title evidence and solicitor confirmation of my entitlement to bring this claim.',
    '2. I seek possession under Ground 1 of Schedule 2 to the Housing Act 1988. I require the Property as my only or principal home.',
    '3. I currently live at 26A Rhodes Place. I rely on my council-tax and utility evidence to confirm that address.',
    '4. My son turns four in December 2026. I intend to move to the Property to support his school application. The intended school is approximately 200 metres from the Property and the Property is in its catchment area.',
    '5. I intend to occupy the Property permanently as my sole and main residence when possession is obtained. I do not intend to use it as a second or holiday home.',
    '6. The tenancy agreement dated 25 November 2025 names the three defendants. I served a new Form 3A notice dated 20 August 2026 and rely only on Ground 1 in that notice.',
    '7. I do not rely on the previous notice, Ground 8 or an allegation that GOA granted the tenancy without my consent. I understand that the court will determine the claim.',
    'Statement of truth',
    'I believe that the facts stated in this witness statement are true.',
    'Signed: ________________________________',
    'Suraraj Pradhan',
    'Dated: _________________________________',
])

standing = evidence / '05_Standing_and_notification_to_tenants'
standing.mkdir(exist_ok=True)
copy2(CASE_SOURCE / 'email-archive-1' / 'STATEMENT OF INTENT AND WITNESS STATEMENT.pdf', evidence / '03_Ground_1' / '03-signed-statement-of-intent-source.pdf')
copy2(CASE_SOURCE / 'email-archive-2' / '2026-03-02-Notifying Ms Shellyann of end of agreement and making aware of rent arrears.pdf', standing / '01-existing-tenant-notification-source.pdf')
copy2(CASE_SOURCE / 'email-archive-2' / '2026-03-30-Handletter to Ms Shelllyann from Mr David for engagemetn.pdf', standing / '02-existing-hand-letter-source.pdf')
copy2(CASE_SOURCE / 'email-archive-2' / '2026-03-30-Attendance at Property on 30 March 2026 and Rent Instructions To Ms Shellanyy from Mr David.pdf', standing / '03-existing-attendance-and-rent-instructions-source.pdf')

(standing / '00-READ-ME-PENDING-PROOF.txt').write_text(
    'These historic documents are retained because they were supplied by the client. They are not treated as proof that every named tenant received a formal GOA-revocation notice. Obtain a formal notice and proof of service on Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne before court issue. Do not rely on the historic arrears assertions for this Ground 1-only claim.',
    encoding='utf-8'
)

(evidence / '00-evidence-index.txt').write_text(
    'EVIDENCE BUNDLE INDEX\n\n'
    '01 Tenancy: signed AST.\n'
    '02 Management and standing: GOA termination notice and management agreement.\n'
    '03 Ground 1: council-tax, utility and supplied signed statement of intent/witness statement.\n'
    '04 Compliance: gas safety certificate, EICR and EPC.\n'
    '05 Standing/tenant notification: historic supplied correspondence retained for review; it is not proof of service on all three tenants.\n\n'
    'To add before issue: official title register, solicitor standing confirmation, formal GOA-revocation notice with proof of service on all tenants, deposit evidence, signed current witness statement, signed Form 3A, N215 and service record.\n',
    encoding='utf-8'
)

(index / '00-IMPORTANT-READ-ME.txt').write_text(
    'GROUND 1 SERVICE PACK - REVISED FOR 20 AUGUST 2026\n\n'
    'Use the editable official Form 3A only after the client has signed it on 20 August 2026. It relies on Ground 1 only. The earliest court date is 20 December 2026.\n\n'
    'This is not court-ready until title/standing, deposit, signed witness statement and actual service evidence are complete. N5, N119 and N215 remain editable official forms; the N5/N119 are pre-issue drafts and the N215 must be signed only after actual service.\n',
    encoding='utf-8'
)

print(ROOT)
