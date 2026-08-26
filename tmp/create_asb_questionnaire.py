from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\t_moh\Documents\GitHub\landlord-heavenv3")
OUTPUT = ROOT / "output" / "client-forms" / "landlord-heaven-serious-antisocial-behaviour-questionnaire.docx"
LOGO = ROOT / "public" / "images" / "logo.png"

DARK = "20103F"
PURPLE = "6D28D9"
LAVENDER = "F5F1FF"
LILAC = "E8E1F8"
TEXT = "3B3550"
MUTED = "675F7D"
RED = "9F1239"
PALE_RED = "FFF1F2"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LILAC, size="8"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for side in ("top", "left", "bottom", "right"):
        tag = "w:" + side
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_width(cell, inches):
    cell.width = Inches(inches)
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(inches * 1440)))
    width.set(qn("w:type"), "dxa")


def fixed_layout(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    properties = table._tbl.tblPr
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_row_height(row, inches):
    properties = row._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(int(inches * 1440)))
    height.set(qn("w:hRule"), "atLeast")
    properties.append(height)


def paragraph(cell_or_document, text="", size=10, bold=False, color=TEXT, align=None, space_after=0, style=None):
    item = cell_or_document.add_paragraph(style=style) if hasattr(cell_or_document, "add_paragraph") else cell_or_document.paragraphs[0]
    item.paragraph_format.space_after = Pt(space_after)
    item.paragraph_format.line_spacing = 1.12
    if align is not None:
        item.alignment = align
    run = item.add_run(text)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return item


def add_field_table(document, fields, widths=(2.2, 4.8)):
    table = document.add_table(rows=0, cols=2)
    fixed_layout(table)
    for label, hint in fields:
        row = table.add_row()
        set_row_height(row, 0.42)
        label_cell, value_cell = row.cells
        set_width(label_cell, widths[0])
        set_width(value_cell, widths[1])
        for cell in (label_cell, value_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
        set_cell_shading(label_cell, LAVENDER)
        paragraph(label_cell, label, size=9, bold=True, color=DARK)
        paragraph(value_cell, hint, size=9, color=MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_section_heading(document, number, title, intro=None):
    table = document.add_table(rows=1, cols=2)
    fixed_layout(table)
    cells = table.rows[0].cells
    set_width(cells[0], 0.55)
    set_width(cells[1], 6.45)
    for cell in cells:
        set_cell_margins(cell, top=95, start=120, bottom=95, end=120)
        set_cell_border(cell, color=PURPLE, size="10")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cells[0], PURPLE)
    set_cell_shading(cells[1], LAVENDER)
    paragraph(cells[0], str(number), size=12, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(cells[1], title, size=13, bold=True, color=DARK)
    if intro:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(intro)
        r.font.name = "Aptos"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)


def add_note_box(document, title, text, fill=LAVENDER, title_color=PURPLE):
    table = document.add_table(rows=1, cols=1)
    fixed_layout(table)
    cell = table.cell(0, 0)
    set_width(cell, 7.0)
    set_cell_margins(cell, top=135, start=160, bottom=135, end=160)
    set_cell_border(cell, color=LILAC, size="10")
    set_cell_shading(cell, fill)
    p = paragraph(cell, title, size=10, bold=True, color=title_color, space_after=3)
    p.paragraph_format.keep_with_next = True
    paragraph(cell, text, size=9.5, color=TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_prompt(document, question, lines=3):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(question)
    r.font.name = "Aptos"
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    table = document.add_table(rows=1, cols=1)
    fixed_layout(table)
    cell = table.cell(0, 0)
    set_width(cell, 7.0)
    set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    set_cell_border(cell)
    set_row_height(table.rows[0], 0.28 * lines)
    paragraph(cell, "Click here and type your answer.", size=9, color=MUTED)


def add_checkbox_line(document, label, options):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    lead = p.add_run(label + "  ")
    lead.font.name = "Aptos"
    lead.font.bold = True
    lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(DARK)
    for index, option in enumerate(options):
        run = p.add_run("☐ " + option + ("     " if index < len(options) - 1 else ""))
        run.font.name = "Segoe UI Symbol"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_evidence_table(document):
    headers = ["Evidence / document", "Held", "Attached", "File name, reference or date", "Notes"]
    widths = [1.75, 0.55, 0.7, 2.1, 1.9]
    rows = [
        "Police crime reference / incident report",
        "Charge, conviction, bail or court document",
        "Criminal behaviour order / IPNA / closure order",
        "Witness statement or incident diary",
        "CCTV, doorbell footage, photographs or video",
        "Tenant, neighbour or management correspondence",
        "Tenancy agreement and relevant tenancy clause",
        "Previous warning letter / breach notice",
        "Council, housing association or environmental-health record",
        "Other relevant material",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    fixed_layout(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        set_width(cell, widths[index])
        set_cell_shading(cell, PURPLE)
        set_cell_border(cell, color=PURPLE)
        set_cell_margins(cell, top=90, start=70, bottom=90, end=70)
        paragraph(cell, text, size=8, bold=True, color=WHITE)
    for label in rows:
        row = table.add_row()
        set_row_height(row, 0.38)
        for index, cell in enumerate(row.cells):
            set_width(cell, widths[index])
            set_cell_border(cell)
            set_cell_margins(cell, top=75, start=70, bottom=75, end=70)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph(row.cells[0], label, size=8.2, color=DARK)
        paragraph(row.cells[1], "☐ Yes  ☐ No", size=8, color=TEXT)
        paragraph(row.cells[2], "☐ Yes  ☐ No", size=8, color=TEXT)
        paragraph(row.cells[3], "", size=8)
        paragraph(row.cells[4], "", size=8)
    return table


def add_page_number(paragraph_obj):
    run = paragraph_obj.add_run("Page ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph_obj._p.append(field)
    run = paragraph_obj.add_run(" of ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "NUMPAGES")
    paragraph_obj._p.append(field)


def add_header_footer(section):
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    fixed_layout(table)
    left, right = table.rows[0].cells
    set_width(left, 4.7)
    set_width(right, 2.3)
    for cell in (left, right):
        set_cell_border(cell, color=WHITE, size="0")
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    if LOGO.exists():
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.95))
    else:
        paragraph(left, "LANDLORD HEAVEN", size=14, bold=True, color=PURPLE)
    paragraph(right, "CASE FACTS QUESTIONNAIRE", size=8, bold=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph(right, "England possession evidence", size=7.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Landlord Heaven  |  Confidential document-preparation questionnaire  |  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(p)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.3)
    add_header_footer(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(9.5)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Serious Antisocial Behaviour & Violence\nLandlord Questionnaire")
    r.font.name = "Aptos Display"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(9)
    r = subtitle.add_run("For fact-gathering before a potential Section 8 possession case in England (Grounds 14, 7A and/or 12).")
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_note_box(
        document,
        "IMMEDIATE SAFETY",
        "If there is an immediate risk to anyone, call 999. Do not confront the tenant, attempt to remove anyone, change locks or take any other self-help eviction step. This form records the facts supplied by the landlord; it is not legal advice and the court decides whether possession is granted.",
        fill=PALE_RED,
        title_color=RED,
    )
    add_note_box(
        document,
        "HOW TO COMPLETE THIS FORM",
        "Type directly into the pale boxes. Give factual information only: say what you personally saw or heard, identify the source of anything reported to you, and attach copies rather than altering originals. Use an extra incident sheet if there is more than one event.",
    )

    add_section_heading(document, 1, "Landlord, property and tenancy details")
    add_field_table(document, [
        ("Landlord full name", ""),
        ("Best email address", ""),
        ("Telephone number", ""),
        ("Correspondence address", ""),
        ("Rental property address", ""),
        ("Letting/managing agent (if any)", ""),
        ("Tenancy start date", ""),
        ("All named tenants on the agreement", ""),
    ])
    add_checkbox_line(document, "Who is currently living at the property?", ["Named tenant(s)", "Family member(s)", "Visitor(s)", "Unknown — explain below"])
    add_prompt(document, "List everyone known to be living at or regularly staying at the property, and explain any uncertainty.", 2)

    add_section_heading(document, 2, "Incident details", "Please complete this section for the most serious incident. Attach extra copies of the next page if there were separate incidents.")
    add_field_table(document, [
        ("Incident date", ""),
        ("Approximate time", ""),
        ("Location", "☐ At the property  ☐ Near the property  ☐ Elsewhere"),
        ("How did you learn about it?", "☐ Saw/heard it myself  ☐ Police  ☐ Witness  ☐ Other"),
        ("Person(s) involved", "Use full names where known; otherwise state how identified."),
        ("Victim / affected person", "Do not include medical details unless necessary and authorised."),
    ])
    add_prompt(document, "Set out a clear factual chronology. Include what happened, who did what, and what you personally observed. Separate known facts from information reported by someone else.", 5)
    add_checkbox_line(document, "Connection to the tenancy or locality", ["At the property", "Affected a neighbour", "Affected the local area", "Connection unclear"])
    add_prompt(document, "Explain why the incident affects the property, people in the locality, or the landlord’s ability to manage the tenancy.", 3)

    document.add_page_break()
    add_section_heading(document, 3, "Police, criminal and court information")
    add_note_box(document, "IMPORTANT", "A serious allegation alone does not establish Ground 7A. If you have a conviction, a relevant order or closure order, provide a readable copy. If the investigation is ongoing, record that accurately rather than guessing the outcome.")
    add_field_table(document, [
        ("Police force / station", ""),
        ("Crime or incident reference", ""),
        ("Officer name / contact (if known)", ""),
        ("Date reported to police", ""),
        ("Arrest made?", "☐ Yes  ☐ No  ☐ Unknown"),
        ("Charge / bail / court date", ""),
        ("Conviction or disposal?", "☐ Yes — attach  ☐ No  ☐ Ongoing / unknown"),
        ("Relevant order?", "☐ Criminal behaviour order  ☐ IPNA  ☐ Closure order  ☐ Other  ☐ None / unknown"),
    ])
    add_prompt(document, "Give the offence, order, hearing date and result (if known). If any bail or other conditions relate to the property or locality, describe them and attach evidence.", 4)
    add_section_heading(document, 4, "Tenancy breach, warnings and current risk")
    add_checkbox_line(document, "Does the tenancy agreement prohibit this conduct or nuisance?", ["Yes", "No", "Not sure — attach agreement"])
    add_checkbox_line(document, "Have warnings or breach letters already been sent?", ["Yes — attach", "No", "Not sure"])
    add_prompt(document, "Summarise any earlier complaints, warnings, incidents or breaches. Include dates, who sent the warning, the tenant’s response and whether the behaviour continued.", 4)
    add_prompt(document, "Describe the current safety or management risk. Identify anyone who may be at risk, and say whether a witness needs their contact details kept confidential from other parties where possible.", 3)

    document.add_page_break()
    add_section_heading(document, 5, "Evidence checklist and document references", "Tick what you hold, attach copies where appropriate, and give file names or reference numbers so the documents can be matched to this questionnaire.")
    add_evidence_table(document)
    add_prompt(document, "Is there any evidence you have not been able to obtain? Explain what it is, who holds it and what steps you have taken.", 3)

    add_section_heading(document, 6, "Notice and case-preparation facts")
    add_field_table(document, [
        ("Address for serving any notice", ""),
        ("All people to be named on a notice", "Use the exact names from the tenancy agreement."),
        ("Has any Section 8 notice been served?", "☐ Yes — attach it and service evidence  ☐ No"),
        ("If served: date and method", ""),
        ("Service method allowed by agreement", "☐ Hand delivery  ☐ Post  ☐ Email  ☐ Not known — attach agreement"),
    ])
    add_checkbox_line(document, "Documents you want considered for preparation", ["Ground 14 notice", "Ground 7A only if qualification is evidenced", "Ground 12 breach ground", "Not sure"])
    add_prompt(document, "List any tenant response, disrepair complaint, safeguarding issue, disability-related issue, or other point that could be raised in a defence. Attach the relevant correspondence.", 4)

    document.add_page_break()
    add_section_heading(document, 7, "Landlord declaration and return checklist")
    add_note_box(document, "BEFORE YOU SIGN", "Check that you have not included assumptions, unverified social-media posts or altered material. Your evidence should be retained in its original form. You remain responsible for deciding whether to serve a notice, for service, filing, and compliance with all legal duties.")
    declaration = [
        "I confirm that the factual information I have provided is true to the best of my knowledge and I have identified where information came from another person.",
        "I confirm that I have authority to provide the documents and contact details supplied with this questionnaire.",
        "I understand this is a document-preparation questionnaire, not legal advice or court representation, and that a court determines any possession claim.",
        "I understand that I must not use self-help eviction, including changing locks or removing a tenant without the lawful process.",
    ]
    for item in declaration:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("☐  " + item)
        r.font.name = "Aptos"
        r.font.size = Pt(9.4)
        r.font.color.rgb = RGBColor.from_string(TEXT)
    add_field_table(document, [
        ("Landlord signature", ""),
        ("Date", ""),
        ("Preferred contact time", ""),
    ])
    add_note_box(document, "RETURN CHECKLIST", "Return this completed questionnaire with: the tenancy agreement; all relevant police/court material; a concise incident chronology; evidence listed in Section 5; prior warnings and tenant correspondence; and any notice or proof of service already used. Keep originals safe and send readable copies only.")
    final = document.add_paragraph()
    final.paragraph_format.space_before = Pt(7)
    final.paragraph_format.space_after = Pt(0)
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run("Landlord Heaven  |  Assisted document preparation for landlords")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
