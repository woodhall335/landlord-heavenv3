from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


SOURCE_ZIP = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-final-pack-20-august-2026-v4.zip")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v7")

if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
with ZipFile(SOURCE_ZIP) as archive:
    archive.extractall(ROOT)

email_path = ROOT / "00_READ_FIRST_CASE_SUMMARY_AND_INDEX" / "06-warm-client-email-to-suraraj.docx"
document = Document(email_path)
body = document._element.body
for paragraph in list(body.findall(qn("w:p"))):
    body.remove(paragraph)

footer = document.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.text = "Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026"
for run in footer.runs:
    run.font.size = Pt(8)

heading = document.add_paragraph()
heading.style = document.styles["Heading 1"]
heading.add_run("Client email - Ground 1 Section 8 Notice and Draft Court Bundle")
for text in [
    "Subject: Ground 1 Section 8 Notice and Draft Court Bundle - 39 Upton Grove",
    "Dear Suraraj,",
    "Thank you for your patience, and apologies for the delay in replying yesterday. I was in meetings and then took time to review the Ground 1 requirements carefully, as I wanted to ensure we were taking a diligent and properly evidenced approach before advising you to proceed.",
    "Please find attached your new Ground 1 Section 8 notice, together with the draft court bundle, service instructions, witness-statement documents and evidence bundle.",
    "The key outstanding document we need from you is an official Land Registry title register/title deed for 39 Upton Grove. This will provide clear evidence of your ownership and materially strengthen the standing aspect of the claim, particularly as GOA Property Solutions Ltd is named within the previous tenancy paperwork.",
    "Please note that the pack includes the termination notice sent to GOA Property Solutions Ltd. If a formal notice confirming GOA's revoked authority was not previously served on all three tenants, we have also included a formal notice for you to sign and serve with the Section 8 notice. Please retain clear proof of service.",
    "The court forms are drafts only and must not be filed until the notice has been correctly served, the service evidence is completed, and the remaining standing evidence has been reviewed.",
    "Kind regards,",
    "Tariq",
    "Landlord Heaven",
]:
    document.add_paragraph(text)
document.save(email_path)

print(ROOT)
