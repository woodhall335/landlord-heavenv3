from docx import Document
from pathlib import Path
for path in [
 Path(r"C:\Users\t_moh\Downloads\will-court-pack\00_READ_FIRST_CASE_SUMMARY_AND_INDEX\04-cover-letter-to-william-and-leanne.docx"),
 Path(r"C:\Users\t_moh\Downloads\will-court-pack\00_READ_FIRST_CASE_SUMMARY_AND_INDEX\03-filing-status-and-next-steps.docx"),
 Path(r"C:\Users\t_moh\Downloads\will-court-pack\01_SERVE_ON_MONDAY_22_JUNE_2026\02-service-checklist-for-monday.docx"),
]:
 print('\n---', path.name, '---')
 doc=Document(path)
 for p in doc.paragraphs:
  if p.text.strip(): print(f'[{p.style.name}] {p.text}')
 print('tables', len(doc.tables))
