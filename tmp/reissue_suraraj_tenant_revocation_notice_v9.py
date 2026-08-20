from pathlib import Path
from shutil import copytree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v8")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v9")
NOTICE = ROOT / "01_SERVE_ON_20_AUGUST_2026" / "05-formal-notice-revocation-of-goa-authority-for-signature.docx"

if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
copytree(SOURCE, ROOT)

document = Document()
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

normal = document.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for style_name, size in [("Title", 15), ("Heading 1", 13)]:
    style = document.styles[style_name]
    style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = True

def add_text(text, bold=False, alignment=None, before=0, after=8):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


add_text("Suraraj Pradhan", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
add_text("26A Rhodes Place\nOldbrook\nMilton Keynes\nMK6 2LX\npradhansuraraj@gmail.com | 07747 817502", alignment=WD_ALIGN_PARAGRAPH.RIGHT, after=14)
add_text("20 August 2026", after=14)
add_text("To: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne", after=0)
add_text("39 Upton Grove\nShenley Lodge\nMilton Keynes\nMK5 7GR", after=16)
add_text("FORMAL NOTICE: REVOCATION OF GOA PROPERTY SOLUTIONS LTD AUTHORITY", bold=True, after=14)
add_text("Dear Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne,", after=10)
add_text("I, Suraraj Pradhan, write personally as the owner of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property).")
add_text("GOA Property Solutions Ltd's authority to act for me in connection with the Property was terminated with effect from 1 April 2026.")
add_text("From that date, GOA Property Solutions Ltd was not authorised to collect rent, vary the tenancy, make commitments or give instructions on my behalf in relation to the Property.")
add_text("This notice does not vary or replace the existing tenancy agreement. It is not a demand for payment of any historic sum. Until you receive a separate signed written notice from me, continue to follow any current written rent-payment arrangement communicated by Black & White. If you are uncertain about a payment instruction, please contact me in writing before making payment.")
add_text("Please send future correspondence about the Property to me at the address above or by email to pradhansuraraj@gmail.com.", after=16)
add_text("Yours sincerely,", after=24)
add_text("Signed: ________________________________", after=0)
add_text("Suraraj Pradhan\nLandlord and issuer of this notice", after=0)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.text = "Formal notice issued personally by Suraraj Pradhan"
for run in footer.runs:
    run.font.name = "Aptos"
    run.font.size = Pt(8)

document.save(NOTICE)
print(ROOT)
