from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


SOURCE = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v12")
DESTINATION = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v13")


def replace_paragraph(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def full_doc_text(document):
    chunks = []

    def collect(part):
        chunks.extend(paragraph.text for paragraph in part.paragraphs)
        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    collect(cell)

    collect(document)
    for section in document.sections:
        collect(section.header)
        collect(section.footer)
    return "\n".join(chunks)


def create_revocation_pdf(path):
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        spaceBefore=5,
        spaceAfter=14,
    )
    signature = ParagraphStyle(
        "Signature",
        parent=body,
        spaceBefore=20,
    )
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
        title="Formal notice of revocation of GOA authority",
        author="Suraraj Pradhan",
    )
    story = [
        Paragraph("Suraraj Pradhan<br/>26A Rhodes Place<br/>Oldbrook<br/>Milton Keynes<br/>MK6 2LX<br/>pradhansuraraj@gmail.com | 07747 817502", body),
        Spacer(1, 8),
        Paragraph("20 August 2026", body),
        Spacer(1, 8),
        Paragraph("To: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne<br/>39 Upton Grove<br/>Shenley Lodge<br/>Milton Keynes<br/>MK5 7GR", body),
        Spacer(1, 8),
        Paragraph("FORMAL NOTICE: REVOCATION OF GOA PROPERTY SOLUTIONS LTD AUTHORITY", heading),
        Paragraph("Dear Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne,", body),
        Paragraph("I, Suraraj Pradhan, write personally as the owner of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property).", body),
        Paragraph("GOA Property Solutions Ltd's authority to act for me in connection with the Property was terminated with effect from 1 April 2026.", body),
        Paragraph("From that date, GOA Property Solutions Ltd was not authorised to collect rent, vary the tenancy, make commitments or give instructions on my behalf in relation to the Property.", body),
        Paragraph("This notice does not vary or replace the existing tenancy agreement. It is not a demand for payment of any historic sum. Until you receive a separate signed written notice from me, continue to follow any current written rent-payment arrangement communicated by Black & White. If you are uncertain about a payment instruction, please contact me in writing before making payment.", body),
        Paragraph("Please send future correspondence about the Property to me at the address above or by email to pradhansuraraj@gmail.com.", body),
        Paragraph("Yours sincerely,", body),
        Paragraph("Signed: ____________________________________<br/>Suraraj Pradhan<br/>Landlord and issuer of this notice", signature),
    ]
    document.build(story)


if DESTINATION.exists():
    shutil.rmtree(DESTINATION)
shutil.copytree(SOURCE, DESTINATION)

service_path = DESTINATION / "01_SERVE_ON_20_AUGUST_2026" / "02-service-instructions-and-checklist.docx"
service_doc = Document(service_path)
for paragraph in service_doc.paragraphs:
    if paragraph.text.startswith("Serve one complete identical set addressed to Shellyann"):
        replace_paragraph(
            paragraph,
            "Prepare THREE complete identical sets: one set addressed to Shellyann Roberts-Henry, one set addressed to Mignal Small and one set addressed to Taylor Goulbourne, each at 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR. Do not rely on one collective set for all three tenants. Each tenant's set must include: (1) the signed Form 3A; (2) the Ground 1 continuation sheet; (3) the signed GOA-revocation notice if it has not already been served on that tenant; and (4) the exact official Renters' Rights Act Information Sheet 2026 PDF and landlord covering letter if it was not previously given to that tenant.",
        )
    if paragraph.text.startswith("Use an independent adult or professional process server"):
        replace_paragraph(
            paragraph,
            "Check the tenancy agreement for any contractual service clause before service and follow it. The preferred method is an independent adult or professional process server delivering each tenant's individually addressed full set. If using post, put each tenant's full set in a separate envelope addressed to that tenant at the Property and retain the posting and tracking evidence. Do not enter, seek confrontation, change locks or attempt self-help eviction.",
        )
    if paragraph.text.startswith("After service, complete the editable N215"):
        replace_paragraph(
            paragraph,
            "After service, complete the editable N215 and the service record using actual facts only, identifying all three individual sets and the actual method used. Suraraj must not sign the N215 statement of truth until satisfied it is accurate. Do not issue a court claim before 20 December 2026.",
        )
service_doc.save(service_path)

record_path = DESTINATION / "01_SERVE_ON_20_AUGUST_2026" / "04-service-record-for-deliverer-complete-immediately.docx"
record_doc = Document(record_path)
for paragraph in record_doc.paragraphs:
    if paragraph.text.startswith("Persons named on the notices:"):
        replace_paragraph(
            paragraph,
            "Three individual complete sets served: one set addressed to Shellyann Roberts-Henry, one set addressed to Mignal Small and one set addressed to Taylor Goulbourne.",
        )
record_doc.save(record_path)

readme = DESTINATION / "01_SERVE_ON_20_AUGUST_2026" / "08-READ-ME-GOA-TERMINATION-NOTICE.txt"
readme.write_text(
    "GOA termination notice already held\n\n"
    "The landlord's 1 April 2026 termination and revocation notice to GOA Property Solutions Ltd is already included in the evidence bundle at:\n"
    "04_EVIDENCE_BUNDLE\\02_Management_and_Standing\\01-termination-notice-to-goa.pdf\n\n"
    "It is not a document to be served on the tenants as part of the Form 3A service today. Keep any email, posting or delivery evidence showing whether and when GOA received it. Do not create or represent a fresh notice as having been served on 1 April 2026.\n",
    encoding="utf-8",
)

pdf_path = DESTINATION / "01_SERVE_ON_20_AUGUST_2026" / "05-formal-notice-revocation-of-goa-authority-for-signature.pdf"
create_revocation_pdf(pdf_path)

revocation_docx = DESTINATION / "01_SERVE_ON_20_AUGUST_2026" / "05-formal-notice-revocation-of-goa-authority-for-signature.docx"
if "Landlord Heaven" in full_doc_text(Document(revocation_docx)):
    raise RuntimeError("Revocation DOCX contains Landlord Heaven branding")
revocation_pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)
if "Landlord Heaven" in revocation_pdf_text:
    raise RuntimeError("Revocation PDF contains Landlord Heaven branding")

print(f"CREATED {DESTINATION}")
print(f"CREATED {pdf_path}")
