from pathlib import Path
from zipfile import ZipFile

from docx import Document


SOURCE_ZIP = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-final-pack-20-august-2026-v5.zip")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v8")

REPLACEMENTS = {
    "William Mumford and Leanne Mumford": "Suraraj Pradhan",
    "39 Boundary Close, Kingswood, Gloucestershire, GL12 8EN": "26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX",
    "William contact": "Client contact",
    "will1883@hotmail.co.uk / 07878 789367": "pradhansuraraj@gmail.com / 07747 817502",
}


def paragraphs_in(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs_in(cell)


def replace_in_paragraph(paragraph):
    original = paragraph.text
    revised = original
    for old, new in REPLACEMENTS.items():
        revised = revised.replace(old, new)
    if revised == original:
        return 0

    for run in paragraph.runs:
        for old, new in REPLACEMENTS.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

    if paragraph.text != revised:
        paragraph.text = revised
    return 1


if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
with ZipFile(SOURCE_ZIP) as archive:
    archive.extractall(ROOT)

updated_documents = 0
updated_paragraphs = 0
for path in sorted(ROOT.rglob("*.docx")):
    document = Document(path)
    changed = 0
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in paragraphs_in(container):
            changed += replace_in_paragraph(paragraph)
    if changed:
        document.save(path)
        updated_documents += 1
        updated_paragraphs += changed

if updated_documents != 15:
    raise SystemExit(f"Expected to update 15 Word documents, updated {updated_documents}")

print(ROOT)
print(f"UPDATED_DOCUMENTS={updated_documents}")
print(f"UPDATED_PARAGRAPHS={updated_paragraphs}")
