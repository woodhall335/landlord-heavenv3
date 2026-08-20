# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copytree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from pypdf import PdfReader, PdfWriter

SOURCE = Path(r'C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v3')
ROOT = Path(r'C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v4')

if ROOT.exists():
    raise SystemExit(f'Remove or rename existing destination before build: {ROOT}')
copytree(SOURCE, ROOT)

index = ROOT / '00_READ_FIRST_CASE_SUMMARY_AND_INDEX'
service = ROOT / '01_SERVE_ON_20_AUGUST_2026'
court = ROOT / '03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'
evidence = ROOT / '04_EVIDENCE_BUNDLE'
template = index / '01-cover-letter-to-suraraj.docx'


def make_doc():
    document = Document(template)
    body = document._element.body
    for paragraph in list(body.findall(qn('w:p'))):
        body.remove(paragraph)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = 'Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026'
    for run in footer.runs:
        run.font.size = Pt(8)
    return document


def write_doc(path, paragraphs):
    document = make_doc()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


write_doc(index / '01-cover-letter-to-suraraj.docx', [
    'Ground 1 possession pack - 39 Upton Grove',
    'Dear Suraraj,',
    'Thank you for confirming that no tenancy deposit was taken. This has been recorded in the pack and removes the deposit-evidence query from the immediate pre-service requirements.',
    'This revised pack is prepared for a new Form 3A notice relying on Ground 1 only. It also contains a separate formal notice confirming that GOA Property Solutions Ltd authority was revoked. If that formal notice was not previously served on every tenant, sign and serve the enclosed notice with the Form 3A on 20 August 2026.',
    'The editable official Form 3A is prepared for service on 20 August 2026. It names Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne and identifies 20 December 2026 as the earliest date for proceedings. You must personally sign it on the actual day of service.',
    'The N5, N119 and N215 are editable official HMCTS forms. The N5 and N119 are pre-issue drafts. The N215 must be checked against what actually occurs, signed only after service and retained with the service record.',
    'This pack is not court-ready until title/standing evidence, the signed witness statement and actual service evidence are complete. The court decides whether Ground 1 is proved.',
    'Yours sincerely',
    'Tariq Mohammed',
    'Landlord Heaven',
])

write_doc(index / '03-what-we-need-from-you-before-service.docx', [
    'What we need from you before Ground 1 service',
    'Dear Suraraj,',
    'Thank you for confirming that no tenancy deposit was taken. We have recorded this in the case file.',
    'Before serving the new Form 3A on 20 August 2026, please send:',
    '1. An official Land Registry title register/title deed for 39 Upton Grove, confirming your ownership.',
    '2. A copy of the formal notice sent to the tenants confirming that GOA Property Solutions Ltd authority had been revoked, together with evidence it was served on Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'If the GOA-revocation notice was not served on all three tenants, please sign and serve the enclosed formal notice now with the Form 3A. Keep an identical copy, complete the service record immediately and complete the N215 only with the actual service facts.',
    'The enclosed Form 3A is dated for service on 20 August 2026 only. Do not serve it late. If these items cannot be checked before service, ask us to prepare a newly dated notice with a revised notice period.',
    'Yours sincerely',
    'Tariq Mohammed',
    'Landlord Heaven',
])

write_doc(index / '06-warm-client-email-to-suraraj.docx', [
    'Subject: Final documents for 20 August service - 39 Upton Grove',
    'Dear Suraraj,',
    'Thank you for confirming that no tenancy deposit was taken. We have updated the file to reflect this.',
    'We have now prepared the revised Ground 1 service pack for 20 August 2026, including the editable official Form 3A, the court-form drafts, a service record and the N215 certificate of service for completion after service.',
    'Before service, please could you send us the following two items:',
    '1. An official Land Registry title register/title deed for 39 Upton Grove, confirming your ownership.',
    '2. A copy of the formal notice sent to the tenants confirming that GOA Property Solutions Ltd authority had been revoked, together with evidence it was served on Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'We have also enclosed a formal GOA-revocation notice for you to sign. If a formal notice was not previously served on all three tenants, please serve this notice with the Form 3A on 20 August and retain the completed service record and proof of delivery.',
    'Once service has taken place, please send us a copy of the signed notices, the completed service record and the delivery evidence. We can then make sure the court-stage material is ready for the next stage.',
    'Kind regards,',
    'Tariq',
    'Landlord Heaven',
])

write_doc(index / '00-case-summary-and-merits-status.docx', [
    'Case summary and merits status',
    'Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR.',
    'Claimant: Suraraj Pradhan. Defendants: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'Route: Ground 1, occupation by landlord or family. The case is deliberately not advanced on Ground 8 because the supplied material records tenant payments to GOA and Black & White and does not safely establish arrears owed by the tenants to Suraraj.',
    'Core case: Suraraj says he will occupy the Property as his sole or principal home. His son turns four in December 2026; the proposed school is about 200 metres away and the Property is within its catchment area. The tenancy began on 25 November 2025. A notice served on 20 August 2026 expires after the first 12 months and gives four months notice.',
    'Deposit status: Suraraj has confirmed that no tenancy deposit was taken. Preserve his written confirmation in the court file.',
    'Status: the legal route is arguable and the documentary case is organised, but it is not ready to issue. Ownership/standing, valid service and a signed current witness statement remain essential conditions.',
])

write_doc(index / '05-merits-risks-and-response-plan.docx', [
    'Ground 1 merits, risks and response plan',
    'Merits: Ground 1 is mandatory if every statutory condition is proved. The proposed 20 August 2026 notice gives four months notice and its earliest proceedings date, 20 December 2026, is after the first 12 months of the tenancy that began on 25 November 2025.',
    'Strengths: all three named tenants are included; the current official Form 3A is used; the case gives a concrete sole/main-home reason; current-address evidence, tenancy evidence and compliance documents are held; no tenancy deposit was taken according to the landlord; and Ground 8 is not pleaded on an unsafe payment theory.',
    'Material risks: the AST names GOA and the management agreement gave GOA letting authority, so Suraraj must retain title evidence and solicitor confirmation that he is the current landlord entitled to sue. The court will scrutinise whether the stated intention to occupy is genuine and still current. Poor service or a notice served late would undermine the claim.',
    'Anticipated response: the tenants may say that GOA was the landlord, that notice was not properly served or that the stated intention is not genuine. The Ground 1 case does not seek a rent judgment. The response is title/standing proof, service evidence, the tenancy and management records, and Suraraj evidence of a genuine continuing intention to make the Property his sole/main home.',
    'Assessment: do not treat this as an 80% to 85% case at this stage. It becomes a credible mandatory Ground 1 claim only once the outstanding evidence is supplied and the notice is validly served. Without those items, it should not be issued.',
])

write_doc(service / '02-service-instructions-and-checklist.docx', [
    'Form 3A and GOA-revocation notice service instructions',
    'SERVICE ON 20 AUGUST 2026 ONLY',
    'Before leaving to serve, check that Suraraj has signed and dated the editable Form 3A and, if it was not previously served on every tenant, the formal GOA-revocation notice. Check that every continuation page is present, the tenant names are correct and the Form 3A states 20 December 2026 as the earliest date for proceedings.',
    'Serve one complete identical set addressed to all three tenants at 39 Upton Grove. The set must include: (1) the signed Form 3A; (2) the Ground 1 continuation sheet; and (3) the signed GOA-revocation notice if it has not already been served on every tenant.',
    'Use an independent adult or professional process server to leave the documents at the Property. Do not enter, do not seek confrontation, do not change locks and do not attempt self-help eviction.',
    'Check the tenancy agreement for any contractual service clause before service. If it requires a specific method, follow it. Keep a copy of every page exactly as served.',
    'Immediately record the actual time, exact manner of delivery, full name and contact details of the deliverer, the address used, a contemporaneous signed service record and any safe photograph of the addressed envelope or delivery. If a postal method is also used, retain the receipt and tracking evidence but do not describe a postal delivery as hand delivery.',
    'After service, complete the editable N215 and the service record using actual facts only. Suraraj must not sign the N215 statement of truth until satisfied it is accurate. Do not issue a court claim before 20 December 2026.',
])

write_doc(service / '04-service-record-for-deliverer-complete-immediately.docx', [
    'Service record - complete immediately after service',
    'Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR.',
    'Documents served: (1) complete Form 3A notice seeking possession (Ground 1); (2) Ground 1 continuation sheet; and (3) formal notice of revocation of GOA authority, if not already validly served on every tenant.',
    'Persons named on the notices: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.',
    'Planned service date: 20 August 2026. Actual date: ____________________. Actual time: ____________________.',
    'Deliverer full name: ____________________. Address/contact number: ____________________.',
    'Exact service method and location: ________________________________________________________________.',
    'Was a copy also posted? Yes / No. If yes, method, receipt/tracking number and posting time: ________________________________________.',
    'Documents and proof retained: signed Form 3A / continuation sheet / GOA-revocation notice if served / photograph if safe / postal proof if used / other: ____________________.',
    'I confirm this record is true and was made immediately after service. Deliverer signature: ____________________. Date: ____________________.',
])

write_doc(service / '05-formal-notice-revocation-of-goa-authority-for-signature.docx', [
    'FORMAL NOTICE: REVOCATION OF GOA PROPERTY SOLUTIONS LTD AUTHORITY',
    'Date: 20 August 2026',
    'To: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne',
    'Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR',
    'I, Suraraj Pradhan, give formal notice that the authority of GOA Property Solutions Ltd to act for me in connection with the Property was terminated with effect from 1 April 2026.',
    'From that date, GOA Property Solutions Ltd was not authorised to collect rent, vary the tenancy, make commitments or give instructions on my behalf in relation to the Property.',
    'This notice does not vary or replace the existing tenancy agreement. It is not a demand for payment of any historic sum. Until you receive a separate signed written notice from me, continue to follow any current written rent-payment arrangement communicated by Black & White. If you are uncertain about a payment instruction, please contact me in writing before making payment.',
    'Please send future correspondence about the Property to me at 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX or by email to pradhansuraraj@gmail.com.',
    'Signed: ________________________________',
    'Suraraj Pradhan',
])

write_doc(court / '01-court-issue-conditions-and-evidence-schedule.docx', [
    'Court issue conditions and evidence schedule',
    'This schedule is for the post-notice court stage. It does not authorise issue.',
    'Do not issue before 20 December 2026. Complete a final review first.',
    'Required Ground 1 evidence: official title evidence; solicitor confirmation of Suraraj standing as landlord; current tenancy agreement; signed Form 3A, continuation sheet, completed N215 and service record; signed current witness statement; current-address evidence; and evidence of genuine continuing intention to occupy as sole/main home.',
    'Deposit position: the landlord confirms no tenancy deposit was taken. Preserve the confirmation in the court file and recheck the position if any contrary evidence emerges.',
    'Required risk checks: any tenant response, defence, repair complaint, vulnerability/equality issue, counterclaim or change in circumstances; and evidence that the former-agent authority/notification point is properly addressed.',
    'Forms: N5 and N119 are editable pre-issue drafts. Court name, claim number, issue date, court fee, statement of truth/signature and service-specific fields must remain blank until the final filing review.',
])

write_doc(ROOT / '05_WITNESS_STATEMENT_FOR_SIGNATURE' / '01-ground-1-witness-statement-for-review-and-signature.docx', [
    'Witness statement of Suraraj Pradhan',
    'IN THE COUNTY COURT AT [TO BE COMPLETED IF CLAIM ISSUED]',
    'Claim No: [TO BE COMPLETED]',
    'Between: Suraraj Pradhan - Claimant',
    'and Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne - Defendants',
    '1. I am Suraraj Pradhan of 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX. I am the owner of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property). I will exhibit official title evidence and solicitor confirmation of my entitlement to bring this claim as Exhibit SP1.',
    '2. I seek possession under Ground 1 of Schedule 2 to the Housing Act 1988. I require the Property as my only or principal home.',
    '3. I currently live at 26A Rhodes Place. I rely on my council-tax and utility evidence as Exhibit SP2 to confirm that address.',
    '4. My son turns four in December 2026. I intend to move to the Property to support his school application. The intended school is approximately 200 metres from the Property and the Property is in its catchment area.',
    '5. I intend to occupy the Property permanently as my sole and main residence when possession is obtained. I do not intend to use it as a second or holiday home.',
    '6. The tenancy agreement dated 25 November 2025 names the three defendants (Exhibit SP3). My management agreement and the termination notice to GOA are Exhibits SP4 and SP5. I served the Form 3A dated 20 August 2026, its Ground 1 continuation sheet and, if required, the GOA-revocation notice, on all three defendants. Before I sign this statement, I will ensure that the actual service date and method are confirmed by the N215 and service record (Exhibit SP6).',
    '7. No tenancy deposit was taken. I do not rely on the previous notice, Ground 8 or an allegation that GOA granted the tenancy without my consent. I understand that the court will determine the claim.',
    'Statement of truth',
    'I believe that the facts stated in this witness statement are true.',
    'Signed: ________________________________',
    'Suraraj Pradhan',
    'Dated: _________________________________',
])

write_doc(evidence / '00-exhibit-schedule-and-bundle-map.docx', [
    'Exhibit schedule and bundle map',
    'SP1 - Official Land Registry title register/title deed and solicitor confirmation of standing. Status: requested before court issue.',
    'SP2 - Current-address evidence: council-tax bill and utility bill. Status: held.',
    'SP3 - Signed assured shorthold tenancy agreement dated 25 November 2025. Status: held.',
    'SP4 - Management agreement with GOA Property Solutions Ltd. Status: held.',
    'SP5 - Termination notice to GOA Property Solutions Ltd. Status: held.',
    'SP6 - Signed Form 3A, Ground 1 continuation sheet, formal GOA-revocation notice if served, N215, service record and delivery proof. Status: to be completed after service.',
    'SP7 - Signed statement of intent/current Ground 1 witness statement and school-related evidence if available. Status: signed source statement held; updated witness statement to be signed before issue.',
    'SP8 - Compliance material: gas safety certificate, EICR and EPC. Status: held. These are labelled background/compliance evidence, not the primary proof of Ground 1.',
    'SP9 - Historic tenant-notification correspondence. Status: held for context only. It is not proof that a formal GOA-revocation notice was served on all three tenants.',
    'Paginate the final court bundle only after SP1 and SP6 have been added. The final witness statement should then refer to the final exhibit/page numbers.',
])

standing = evidence / '05_Standing_and_notification_to_tenants'
(standing / '00-READ-ME-PENDING-PROOF.txt').write_text(
    'These historic documents are retained because they were supplied by the client. They are not treated as proof that every named tenant received a formal GOA-revocation notice. A new formal notice is in the service folder. If no valid prior notice was served on Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne, sign and serve it with the Form 3A on 20 August 2026, then retain service proof. Do not rely on historic arrears assertions for this Ground 1-only claim.',
    encoding='utf-8'
)
(evidence / '00-evidence-index.txt').write_text(
    'EVIDENCE BUNDLE INDEX\n\n'
    '01 Tenancy: signed AST.\n'
    '02 Management and standing: GOA termination notice and management agreement.\n'
    '03 Ground 1: council-tax, utility and supplied signed statement of intent/witness statement.\n'
    '04 Compliance: gas safety certificate, EICR and EPC.\n'
    '05 Standing/tenant notification: historic supplied correspondence retained for context; it is not proof of service on all three tenants.\n'
    '06 Exhibit schedule: exhibit map and court-bundle pagination instructions.\n\n'
    'To add before issue: official title register, solicitor standing confirmation, signed Form 3A and continuation sheet, N215, service record, proof of service, formal GOA-revocation notice if served and signed current witness statement.\n',
    encoding='utf-8'
)

n215 = ROOT / '02_COMPLETE_AFTER_SERVICE' / '01-n215-ground-1-service-certificate-editable-complete-after-service.pdf'
reader = PdfReader(str(n215))
writer = PdfWriter()
writer.clone_document_from_reader(reader)
writer.reattach_fields()
for page in writer.pages:
    writer.update_page_form_field_values(page, {
        'Text1': 'Form 3A notice seeking possession (Ground 1), Ground 1 continuation sheet and, if served, the formal GOA-revocation notice.',
    }, auto_regenerate=True)
with open(n215, 'wb') as output:
    writer.write(output)

print(ROOT)
