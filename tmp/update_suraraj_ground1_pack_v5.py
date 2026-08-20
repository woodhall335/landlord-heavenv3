from pathlib import Path
from shutil import copytree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, NumberObject, TextStringObject


SOURCE = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v4")
ROOT = Path(r"C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-20-august-2026-final-v5")


if ROOT.exists():
    raise SystemExit(f"Destination already exists: {ROOT}")
copytree(SOURCE, ROOT)

INDEX = ROOT / "00_READ_FIRST_CASE_SUMMARY_AND_INDEX"
SERVICE = ROOT / "01_SERVE_ON_20_AUGUST_2026"
COURT = ROOT / "03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES"
EVIDENCE = ROOT / "04_EVIDENCE_BUNDLE"
WITNESS = ROOT / "05_WITNESS_STATEMENT_FOR_SIGNATURE"
TEMPLATE = INDEX / "01-cover-letter-to-suraraj.docx"


def make_document(footer="Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026"):
    document = Document(TEMPLATE)
    body = document._element.body
    for paragraph in list(body.findall(qn("w:p"))):
        body.remove(paragraph)
    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.text = footer
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)
    return document


def save_document(path, title, paragraphs, footer=None):
    document = make_document(footer or "Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026")
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 1"]
    heading.add_run(title)
    for text in paragraphs:
        if text.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Heading 2"]
            paragraph.add_run(text[3:])
        elif text.startswith("• "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(text[2:])
        else:
            document.add_paragraph(text)
    document.save(path)


def configure_multiline_fields(writer, field_names, font_size=8):
    for page in writer.pages:
        for annotation in page.get("/Annots", []):
            widget = annotation.get_object()
            if widget.get("/T") in field_names:
                widget[NameObject("/Ff")] = NumberObject(int(widget.get("/Ff", 0)) | 4096)
                widget[NameObject("/DA")] = TextStringObject(f"/ArialMT {font_size} Tf 0 g")


def update_n119():
    path = COURT / "03-draft-n119-ground-1-pre-issue-editable.pdf"
    reader = PdfReader(str(path))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    writer.reattach_fields()
    configure_multiline_fields(writer, {"Text Field 109", "Text Field 113", "Text Field 114", "Text Field 117"}, 8)
    values = {
        "Text Field 109": (
            "Written assured shorthold tenancy agreement dated 25 November 2025, "
            "which became an assured periodic tenancy on 1 May 2026."
        ),
        "Text Field 113": (
            "The claimant seeks possession under Ground 1 of Schedule 2 to the Housing Act 1988 only. "
            "He genuinely requires the Property as his only or principal home and intends to move there with "
            "his family. His son turns four in December 2026; the Property is about 200 metres from the school "
            "relevant to the planned application. The claimant relies on title/standing evidence, current-address "
            "evidence, a signed witness statement and service evidence. The written tenancy agreement identifies "
            "GOA Property Solutions Ltd as landlord; the claimant's standing case is set out in his signed statement "
            "and supporting documents. He does not contend that revocation of GOA's authority alone terminated or "
            "varied the tenancy. Form 3A and its continuation sheet were served on each defendant on 20 August 2026. "
            "The notice did not permit proceedings before 20 December 2026. No Ground 8, rent-arrears or other "
            "money claim is made."
        ),
        "Text Field 114": "",
        "Text Field 115": "Form 3A notice seeking possession",
        "Text Field 103a": "20082026",
        "Text Field 117": "Not applicable - no separate financial or other information is relied upon under section 8.1.",
        "Check Box 27": "/Yes",
        "Check Box 51": "/Yes",
    }
    for page in writer.pages:
        writer.update_page_form_field_values(page, values, auto_regenerate=True)
    with open(path, "wb") as output:
        writer.write(output)


save_document(
    INDEX / "00-case-summary-and-merits-status.docx",
    "Case summary and merits status",
    [
        "Property: 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR.",
        "Claimant: Suraraj Pradhan. Defendants: Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne.",
        "Route: Ground 1, occupation by landlord or family. This is a Ground 1-only case. It does not rely on Ground 8, rent arrears or a money claim because the supplied material records tenant payments to GOA and Black & White and does not safely establish arrears owed by the tenants to Suraraj.",
        "Tenancy status: the written agreement is dated 25 November 2025 and records an initial assured shorthold tenancy. It became an assured periodic tenancy on 1 May 2026. The 12-month period is relevant to Ground 1 timing, but the original fixed end date is not relied upon.",
        "Core case: Suraraj says he will occupy the Property as his sole or principal home. His son turns four in December 2026; the proposed school is about 200 metres away and the Property is within its catchment area. A notice served on 20 August 2026 gives four months' notice and does not permit proceedings before 20 December 2026.",
        "Status: this is a prepared service pack, not a court-ready claim. Standing must be resolved by evidence and, because the written agreement identifies GOA as landlord, independent legal advice before issue is strongly recommended. Actual service proof, title evidence, a signed current witness statement and confirmed deposit position also remain essential.",
    ],
)

save_document(
    INDEX / "02-filing-status-and-next-steps.docx",
    "Filing status and next steps",
    [
        "## Current status",
        "A new editable Form 3A is prepared on Ground 1 only for service on 20 August 2026. The old notice is not relied upon. The earliest date for proceedings is 20 December 2026.",
        "## Before service on 20 August 2026",
        "1. Check the tenant names, property address and landlord contact details against the written tenancy agreement.",
        "2. Obtain the official title register and the documents showing GOA's role and the basis on which Suraraj is the proper claimant. The AST naming GOA creates a standing risk which a self-declaration does not cure.",
        "3. Keep Suraraj's written confirmation that no tenancy deposit was taken. Obtain GOA's written confirmation too, if possible. If any deposit was received, obtain scheme and prescribed-information evidence and take legal advice before issue.",
        "4. Confirm whether the required Renters' Rights Act information sheet was served by 31 May 2026, and retain the evidence or obtain legal advice about any non-compliance.",
        "5. Suraraj must sign and date the Form 3A on 20 August 2026. Serve a full identical copy on all three tenants and complete the service record immediately.",
        "## After service",
        "Complete the N215 only with actual service facts. Keep the signed Form 3A, every continuation page, the formal GOA-revocation notice if served, the service record, delivery proof and any tenant response.",
        "## Before court issue",
        "Do not issue before 20 December 2026. Obtain independent legal advice on standing, then recheck the continuing intention to occupy, title evidence, deposit position, witness statement, compliance material, service proof, tenant correspondence, repair issues, vulnerability/equality issues and any counterclaim.",
    ],
)

save_document(
    INDEX / "04-bundle-index-and-document-status.docx",
    "Bundle index and document status",
    [
        "1. Read first: case summary, cover letter, next steps, evidence request and merits/risk review.",
        "2. Serve on 20 August 2026: editable Form 3A, statutory continuation sheet, service instructions, service record, GOA-revocation notice and do-not-serve warning.",
        "3. Complete after service: editable N215 certificate of service. It must reflect actual service and is not to be signed beforehand.",
        "4. Court stage after 20 December 2026: editable N5 and N119 drafts, court-issue schedule and litigant-in-person hearing guide. They are not to be signed or filed before the final review.",
        "5. Evidence bundle: tenancy, GOA management/termination, current-address, compliance, signed Ground 1 source statement and historic tenant-notification correspondence retained for review.",
        "6. Witness evidence: claimant's standing statement and updated Ground 1 witness statement, both for Suraraj to review, sign and date shortly before issue.",
    ],
)

save_document(
    INDEX / "05-merits-risks-and-response-plan.docx",
    "Ground 1 merits, risks and response plan",
    [
        "Merits: Ground 1 is mandatory only if every statutory condition is proved. The proposed 20 August 2026 notice gives four months' notice and its earliest proceedings date, 20 December 2026, is after the first 12 months of the tenancy that began on 25 November 2025.",
        "Strengths: all three named tenants are included; the current official Form 3A is used; the case gives a concrete sole/main-home reason; current-address evidence and compliance documents are held; and Ground 8 is not pleaded on an unsafe payment theory.",
        "Material risk - standing: the written tenancy agreement identifies GOA Property Solutions Ltd as landlord. The Land Registry title proves ownership but does not, by itself, prove that Suraraj is the current landlord entitled to bring this claim. Terminating GOA's management authority does not by itself terminate or assign the tenancy. Obtain independent legal advice on this issue before filing.",
        "Other risks: the court will scrutinise whether the stated intention to occupy is genuine and still current. Poor service, a defective notice, an undisclosed deposit, tenant vulnerability, disrepair or a counterclaim could delay or defeat the claim.",
        "Response plan: keep the case Ground 1-only; prove service and title; present the management/standing documents accurately; do not advance historic arrears as a claim; and obtain legal advice rather than making unsupported submissions about GOA's legal status.",
        "Assessment: do not treat this as an 80% to 85% case. It may become a credible Ground 1 claim only once standing is independently verified, the outstanding evidence is supplied and the notice has been validly served. Without those items, it should not be issued.",
    ],
)

save_document(
    COURT / "01-court-issue-conditions-and-evidence-schedule.docx",
    "Court issue conditions and evidence schedule",
    [
        "This schedule is for the post-notice court stage. It does not authorise issue.",
        "Do not issue before 20 December 2026. Complete a final review first.",
        "Required Ground 1 evidence: official title register; evidence and independent advice confirming that Suraraj is entitled to sue despite the AST naming GOA; current tenancy agreement; signed Form 3A and continuation sheet; completed N215 and service record; signed current witness statement; current-address evidence; and evidence of genuine continuing intention to occupy as sole/main home.",
        "Deposit position: Suraraj confirms that no tenancy deposit was taken. Preserve that confirmation and seek confirmation from GOA if possible. If contrary evidence emerges, stop and take legal advice before issue.",
        "Compliance check: confirm whether the Renters' Rights Act information sheet was served by 31 May 2026 and retain service evidence or take legal advice about any failure.",
        "Required risk checks: any tenant response, defence, repair complaint, vulnerability/equality issue, counterclaim or change in circumstances; and evidence that the former-agent authority/notification point is properly addressed.",
        "Forms: N5 and N119 are editable pre-issue drafts. Court name, claim number, issue date, court fee and statements of truth/signatures must remain blank until the final filing review. The N119 uses the correct assured-periodic-tenancy description.",
    ],
)

save_document(
    EVIDENCE / "00-exhibit-schedule-and-bundle-map.docx",
    "Exhibit schedule and bundle map",
    [
        "SP1 - Official Land Registry title register/title deed and signed claimant's statement of standing and ownership. Status: title register requested; statement ready for signature. Independent standing advice remains required before issue because the AST identifies GOA as landlord.",
        "SP2 - Current-address evidence: council-tax bill and utility bill. Status: held.",
        "SP3 - Written tenancy agreement dated 25 November 2025. Status: held. It records the original AST terms; the tenancy became assured periodic on 1 May 2026.",
        "SP4 - Management agreement with GOA Property Solutions Ltd. Status: held.",
        "SP5 - Termination notice to GOA Property Solutions Ltd. Status: held.",
        "SP6 - Signed Form 3A, Ground 1 continuation sheet, formal GOA-revocation notice if served, N215, service record and delivery proof. Status: complete after actual service only.",
        "SP7 - Signed statement of intent and current Ground 1 witness statement; school-related evidence if available. Status: source statement held; updated witness statement ready for signing before issue.",
        "SP8 - Compliance material: gas safety certificate, EICR and EPC. Status: held. These are labelled background/compliance evidence, not the primary proof of Ground 1.",
        "SP9 - Historic tenant-notification correspondence. Status: held for context only. It is not proof that a formal GOA-revocation notice was served on all three tenants.",
        "Paginate the final court bundle only after SP1 and SP6 have been added. Update exhibit/page references in the final signed witness statement after pagination.",
    ],
)

save_document(
    WITNESS / "00-claimants-statement-of-standing-and-ownership-for-signature.docx",
    "Claimant's statement of standing and ownership",
    [
        "IN THE COUNTY COURT AT [TO BE COMPLETED IF CLAIM ISSUED]",
        "Claim No: [TO BE COMPLETED]",
        "Between: Suraraj Pradhan - Claimant",
        "and Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne - Defendants",
        "1. I, Suraraj Pradhan, of 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX, am the registered freehold proprietor of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property).",
        "2. I attach an official copy of the Land Registry title register for the Property at exhibit SP1.",
        "3. GOA Property Solutions Ltd was engaged in connection with the management of the Property. I rely on the management agreement dated 1 May 2025 at SP4 and the written tenancy agreement dated 25 November 2025 at SP3.",
        "4. On 1 April 2026, I terminated GOA Property Solutions Ltd's authority to manage the Property. The termination notice is at SP5.",
        "5. On 20 August 2026, the Form 3A notice and the notice explaining the revocation of GOA's authority were served on each Defendant. The actual service evidence is at SP6.",
        "6. I contend that, as registered owner and the person entitled to the landlord's interest, I am the proper claimant to bring these possession proceedings. I rely on the title register and the documents identified above.",
        "7. I do not contend that the revocation notice alone terminated, assigned or varied the tenancy. It was served to notify the Defendants of GOA's revoked authority and my position as claimant.",
        "Statement of Truth",
        "I believe that the facts stated in this statement are true.",
        "Signed: ________________________________",
        "Suraraj Pradhan",
        "Dated: _________________________________",
    ],
)

save_document(
    WITNESS / "01-ground-1-witness-statement-for-review-and-signature.docx",
    "Witness statement of Suraraj Pradhan",
    [
        "IN THE COUNTY COURT AT [TO BE COMPLETED IF CLAIM ISSUED]",
        "Claim No: [TO BE COMPLETED]",
        "Between: Suraraj Pradhan - Claimant",
        "and Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne - Defendants",
        "1. I am Suraraj Pradhan of 26A Rhodes Place, Oldbrook, Milton Keynes, MK6 2LX. I am the Claimant.",
        "2. I seek possession of 39 Upton Grove, Shenley Lodge, Milton Keynes, MK5 7GR (the Property) under Ground 1 of Schedule 2 to the Housing Act 1988 only.",
        "3. I am the registered freehold proprietor of the Property. I rely on the official title register and my signed statement of standing and ownership at SP1.",
        "4. The written tenancy agreement is dated 25 November 2025. It records rent of £1,700 per calendar month and identifies GOA Property Solutions Ltd as landlord. A copy is at SP3. The tenancy became an assured periodic tenancy on 1 May 2026.",
        "5. GOA Property Solutions Ltd had been engaged in connection with management of the Property. I terminated GOA's management authority on 1 April 2026. The management agreement and termination notice are at SP4 and SP5.",
        "6. On 20 August 2026, I served each Defendant with Form 3A dated 20 August 2026, the Ground 1 continuation sheet containing the statutory wording and reasons, and the notice confirming that GOA's authority had been revoked. Service was effected by [INSERT ACTUAL METHOD]. The completed N215, service record and supporting proof are at SP6.",
        "7. The Form 3A did not permit proceedings before 20 December 2026. I do not rely on an earlier notice.",
        "8. No tenancy deposit was paid or received in connection with this tenancy, to the best of my knowledge and belief. I have not received or held a tenancy deposit for the Defendants.",
        "9. This claim does not rely on Ground 8, rent arrears or any money claim against the Defendants.",
        "10. I genuinely require the Property as my only or principal home. I intend to move there with my family once possession is recovered and do not intend to use it as a second or holiday home.",
        "11. My son turns four in December 2026. The Property is approximately 200 metres from the school relevant to his planned school application. This supports the practical and genuine need for my family to live at the Property.",
        "12. I rely on my council-tax and utility evidence at SP2, my signed statement of intention at SP7, and the documents identified above. I respectfully ask the Court to make a possession order.",
        "Statement of Truth",
        "I believe that the facts stated in this witness statement are true. I understand that proceedings for contempt of court may be brought against anyone who makes, or causes to be made, a false statement in a document verified by a statement of truth without an honest belief in its truth.",
        "Signed: ________________________________",
        "Suraraj Pradhan",
        "Dated: _________________________________",
    ],
)

save_document(
    COURT / "04-litigant-in-person-hearing-guide.docx",
    "Litigant-in-person hearing guide",
    [
        "## Bring",
        "• Issued N5 and N119, the signed witness statement, the claimant's standing statement, the paginated exhibit bundle, Form 3A and continuation sheet, N215 and service proof, and the official title register.",
        "• A short note of the order requested and a copy for the Defendants if the court directs.",
        "## What to say",
        "Good morning, Judge. I am Suraraj Pradhan, the claimant and a litigant in person. This is a Ground 1 possession claim only. I seek possession because I genuinely intend to make the Property my only or principal home with my family.",
        "The documents show: first, my ownership and the standing evidence; second, service on all three Defendants; third, that the notice did not permit proceedings before 20 December 2026; and fourth, my genuine continuing intention to occupy the Property. I do not seek rent arrears or make a Ground 8 claim.",
        "## If GOA is raised",
        "The written tenancy agreement identifies GOA Property Solutions Ltd. My position on standing is set out in my signed statement and supported by the title register, management agreement and termination evidence. I do not say that revoking GOA's authority alone terminated or varied the tenancy. I ask the Court to determine standing on the evidence before it.",
        "Do not speculate, make unsupported submissions, argue about historic rent, or say that tenants owe arrears. If the judge identifies a standing defect, ask for directions or an adjournment to obtain legal advice.",
        "## Court conduct",
        "Arrive early, address the judge as Sir or Madam unless told otherwise, speak only to the judge, take the court to the relevant exhibit/page, answer questions directly and keep a note of every order or direction.",
    ],
)

(EVIDENCE / "00-evidence-index.txt").write_text(
    "EVIDENCE BUNDLE INDEX\n\n"
    "SP1 Standing and ownership: official title register and signed claimant's standing statement. Independent standing advice remains required before issue.\n"
    "SP2 Current address: council-tax and utility evidence.\n"
    "SP3 Tenancy: written agreement dated 25 November 2025. It records the original AST terms; it became assured periodic on 1 May 2026.\n"
    "SP4 Management: GOA management agreement.\n"
    "SP5 GOA termination: termination notice to GOA.\n"
    "SP6 Service: signed Form 3A, continuation sheet, formal GOA-revocation notice if served, N215, service record and delivery proof.\n"
    "SP7 Ground 1 evidence: signed statement of intent, current witness statement and school-related evidence if available.\n"
    "SP8 Compliance: gas safety certificate, EICR and EPC. These are background/compliance evidence, not primary Ground 1 proof.\n"
    "SP9 Historic tenant-notification correspondence: context only, not proof of a formal notice served on all tenants.\n\n"
    "Do not issue until the title register, standing evidence/advice, actual service proof, signed witness evidence, deposit confirmation and Renters' Rights Act information-sheet compliance check have been completed.\n",
    encoding="utf-8",
)

(EVIDENCE / "05_Standing_and_notification_to_tenants" / "00-READ-ME-PENDING-PROOF.txt").write_text(
    "These historic documents are retained because they were supplied by the client. They are not proof that every named tenant received a formal GOA-revocation notice. The signed claimant's standing statement is evidence only; it does not replace independent legal advice about standing where the written AST identifies GOA as landlord. Obtain the official title register, standing advice and actual service proof before court issue. Do not rely on historic arrears assertions for this Ground 1-only claim.",
    encoding="utf-8",
)

(INDEX / "00-IMPORTANT-READ-ME.txt").write_text(
    "GROUND 1 SERVICE PACK - REVISED FOR 20 AUGUST 2026\n\n"
    "Use the editable official Form 3A only after Suraraj has signed it on 20 August 2026. It relies on Ground 1 only. The earliest court date is 20 December 2026.\n\n"
    "This is a service pack, not a court-ready claim. The written tenancy agreement identifies GOA as landlord, so standing needs independent legal advice before issue. N5, N119 and N215 remain editable official forms; N5/N119 are pre-issue drafts and N215 must be signed only after actual service.\n",
    encoding="utf-8",
)

update_n119()

print(ROOT)
