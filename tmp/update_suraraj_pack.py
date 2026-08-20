# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from pypdf import PdfReader, PdfWriter
ROOT=Path(r'C:\Users\t_moh\Downloads\suraraj-pradhan-ground-1-client-pack-19-august-2026')
FORMS=Path(r'C:\Users\t_moh\Documents\GitHub\landlord-heavenv3\tmp\suraraj-authoring-sources')
TEMPLATE=ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'01-cover-letter-to-suraraj.docx'

def make_doc(out):
    doc=Document(TEMPLATE)
    body=doc._element.body
    for p in list(body.findall(qn('w:p'))): body.remove(p)
    footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.text='Landlord Heaven | 39 Upton Grove possession pack | Prepared 19 August 2026'
    for r in footer.runs:r.font.size=Pt(8)
    return doc

def p(doc,t,style='Normal'): doc.add_paragraph(t,style=style)

doc=make_doc(ROOT/'x')
p(doc,'What we need from you to finalise the Ground 1 notice')
p(doc,'Dear Suraraj,')
p(doc,'Thank you for confirming the Ground 1 route. Before the new notice is served, please send the following two documents:')
p(doc,'1. Official Land Registry title register/title deed for 39 Upton Grove. This is required to evidence your ownership and should be kept for the court-stage bundle.')
p(doc,"2. A copy of the formal written notice sent to the tenants confirming GOA Property Solutions Ltd's authority had been revoked, together with evidence of service on each of Shellyann Roberts-Henry, Mignal Small and Taylor Goulbourne. The existing correspondence in the file is not evidence of notice to all three named tenants.")
p(doc,'Please also confirm in writing that no tenancy deposit was taken from any tenant. If a deposit was taken, send the amount, date received, recipient, scheme details and prescribed information before the notice is served.')
p(doc,'Once received and checked, we will retain the documents in the evidence bundle. The notice must be signed by you and served with all pages on each tenant. The court-issue forms are prepared in draft only and must not be signed or filed until after the notice has expired and the service evidence is complete.')
p(doc,'Important: the enclosed Form 3A is dated 19 August 2026 and is for service on that date only. If the required documents are not received and checked in time, do not serve it late; ask us to issue a newly dated notice with a revised four-month notice period.')
p(doc,'Yours sincerely\nTariq Mohammed\nLandlord Heaven')
doc.save(ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'03-what-we-need-from-you-before-service.docx')

# update simple readme with extra requirements
readme=ROOT/'00_READ_FIRST_CASE_SUMMARY_AND_INDEX'/'00-IMPORTANT-READ-ME.txt'
requirement = 'Before service, obtain: (1) official Land Registry title evidence; (2) formal GOA-revocation notice served on every tenant; and (3) written confirmation that no tenancy deposit was taken.'
existing_readme = readme.read_text(encoding='utf-8')
base_readme = existing_readme.replace(requirement, '').strip()
readme.write_text(base_readme+'\n\n'+requirement+'\n',encoding='utf-8')

def make_draft_form(source, destination, values, static_draw=None):
    reader = PdfReader(str(source))
    if static_draw:
        first_box = reader.pages[0].mediabox
        page_size = (float(first_box.width), float(first_box.height))
        static_path = ROOT / 'tmp-static-n5.pdf'
        static = canvas.Canvas(str(static_path), pagesize=page_size)
        for page_number in range(len(reader.pages)):
            static_draw(static, page_number)
            static.showPage()
        static.save()
        static_reader = PdfReader(str(static_path))
        merged = PdfWriter()
        for index, page in enumerate(reader.pages):
            page.merge_page(static_reader.pages[index])
            merged.add_page(page)
        prepared = ROOT / 'tmp-prepared-form.pdf'
        with open(prepared, 'wb') as output:
            merged.write(output)
        reader = PdfReader(str(prepared))
        prepared.unlink(missing_ok=True)
        static_path.unlink(missing_ok=True)
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    writer.reattach_fields()
    for page in writer.pages:
        writer.update_page_form_field_values(page, values, auto_regenerate=True)
    with open(destination, 'wb') as output:
        writer.write(output)

tenant_names = 'Shellyann Roberts-Henry\nMignal Small\nTaylor Goulbourne'
property_address = '39 Upton Grove\nShenley Lodge\nMilton Keynes\nMK5 7GR'

n5 = ROOT/'03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'/'02-draft-n5-ground-1-pre-issue.pdf'
def draw_n5_static(page, page_number):
    if page_number != 0:
        return
    page.setFont('Helvetica', 8)
    for index, name in enumerate(tenant_names.split('\n')):
        page.drawString(31, 584 - index * 9, name)
    for index, line in enumerate(property_address.split('\n')):
        page.drawString(31, 526 - index * 9, line)
        page.drawString(31, 430 - index * 9, line)
    page.setFont('Helvetica-Bold', 10)
    page.drawString(31, 315, 'X')
    page.drawString(31, 254, 'X')
make_draft_form(FORMS/'N5_0526.pdf', n5, {
}, draw_n5_static)

n119 = ROOT/'03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES'/'03-draft-n119-ground-1-pre-issue.pdf'
def draw_n119_static(page, page_number):
    page.setFont('Helvetica-Bold', 10)
    if page_number == 1:
        page.drawString(59, 470, 'X')
    elif page_number == 3:
        page.drawString(59, 452, 'X')
make_draft_form(FORMS/'N119_0526.pdf', n119, {
    'Text Field 3': 'Suraraj Pradhan',
    'Text Field 4': tenant_names,
    'Text Field 100': '39 Upton Grove',
    'Text Field 99': 'Shenley Lodge',
    'Text Field 98': 'Milton Keynes',
    'Text Field 96': 'MK5 7GR',
    'Text Field 108': tenant_names,
    'Text Field 109': 'Assured shorthold tenancy (original written AST dated 25 November 2025)',
    'Text Field 102': '25112025',
    'Text Field 110': '1,700',
    'Text Field 113': ('The claimant seeks possession under Ground 1 of Schedule 2 to the Housing Act 1988. '
                       'The claimant requires the property as his only or principal home. The tenancy began on '
                       '25 November 2025. A new Form 3A dated 19 August 2026 is included for service on all three '
                       'defendants. Before issue, this draft must be updated with the actual service evidence. The '
                       'earliest date for proceedings is 19 December 2026. The claimant relies on the tenancy '
                       'agreement, official title evidence, current-address evidence, service evidence and his signed '
                       'witness statement. The claimant does not rely on Ground 8.'),
    'Text Field 114': 'No rent arrears are claimed in this Ground 1-only draft.',
    'Text Field 115': 'Form 3A notice seeking possession',
}, draw_n119_static)

for superseded in (
    ROOT / 'tmp-form3a-overlay.pdf',
    ROOT / '03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES' / '02-current-official-n5-england-05-26.pdf',
    ROOT / '03_COURT_FORMS_FOR_ISSUE_AFTER_NOTICE_EXPIRES' / '03-current-official-n119-england-05-26.pdf',
):
    superseded.unlink(missing_ok=True)
print('updated')
